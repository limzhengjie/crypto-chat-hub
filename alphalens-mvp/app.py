"""
AlphaLens — AI-Powered Crypto Research Agent
─────────────────────────────────────────────
Dashboard : Binance WebSocket → SQLite → Plotly charts
Chat      : Ask anything → GPT-4o agent fetches from CoinGecko, DefiLlama, Binance → cited answer
Deep Dive : One-click comprehensive report from all data sources
"""

import html
import io
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from dotenv import load_dotenv
from docx import Document

from src.database import init_db
from src.indicators import add_indicators, indicator_summary
from src.metric_tooltip import render_metric_with_tooltip
from src.ws_client import BinanceKlineStream
from src.orderbook import BinanceOrderBookStream
from src.history import fetch_historical_klines
from src.agent import run_agent

load_dotenv()

# Make Streamlit Community Cloud secrets available as env vars
try:
    for key, value in st.secrets.items():
        if isinstance(value, str):
            os.environ.setdefault(key, value)
except FileNotFoundError:
    pass  # No secrets.toml locally — rely on .env via load_dotenv()

AVAILABLE_SYMBOLS = [
    "BTCUSDT",
    "ETHUSDT",
    "SOLUSDT",
    "BNBUSDT",
    "XRPUSDT",
    "DOGEUSDT",
    "ADAUSDT",
    "AVAXUSDT",
]

CHART_COLORS = ["#7c83fd", "#00e676", "#ff9800", "#e91e63"]

TOOL_LABELS = {
    "get_market_data": "📊 Fetching market data from CoinGecko",
    "get_tvl": "🔗 Fetching TVL from DefiLlama",
    "get_funding_rate": "📈 Fetching funding rates from Binance Futures",
    "get_open_interest": "📉 Fetching open interest from Binance Futures",
    "get_technical_analysis": "🔬 Running technical analysis on Binance data",
    "get_prediction_markets": "🎯 Fetching prediction markets from Polymarket",
    "get_prediction_accuracy": "📊 Checking Polymarket accuracy data",
    "get_crypto_news": "📰 Fetching latest news from RSS feeds",
}

DEEP_DIVE_TOKENS = [
    "BTC",
    "ETH",
    "SOL",
    "BNB",
    "XRP",
    "DOGE",
    "ADA",
    "AVAX",
    "DOT",
    "LINK",
    "UNI",
    "NEAR",
    "ARB",
    "OP",
]

SCANNER_COIN_OPTIONS = [
    "All",
    "BTC",
    "ETH",
    "SOL",
    "BNB",
    "XRP",
    "DOGE",
    "ADA",
    "AVAX",
    "LINK",
    "DOT",
    "UNI",
    "NEAR",
    "ARB",
    "SUI",
    "APT",
    "PEPE",
    "TON",
    "TIA",
]

SCANNER_TIME_OPTIONS = ["All", "Today", "This Week", "This Month", "Long-term"]
_DASH_INTERVAL_OPTS = ["1m", "3m", "5m"]


def _report_to_docx_bytes(report_text: str, symbol: str) -> bytes:
    """Convert markdown-like report text to a .docx file in memory."""

    def _add_markdown_runs(paragraph, text: str) -> None:
        """
        Render simple markdown emphasis to Word runs.
        Supports: **bold**, __bold__, *italic*, _italic_.
        """
        pattern = r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)"
        idx = 0
        for match in re.finditer(pattern, text):
            start, end = match.span()
            if start > idx:
                paragraph.add_run(text[idx:start])

            token = match.group(0)
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("__") and token.endswith("__"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("_") and token.endswith("_"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            idx = end

        if idx < len(text):
            paragraph.add_run(text[idx:])

    doc = Document()
    doc.add_heading(f"AlphaLens Deep Dive Report - {symbol}", level=0)

    for raw_line in report_text.splitlines():
        line = raw_line.strip()

        if not line:
            # Keep section spacing in the generated document.
            doc.add_paragraph("")
            continue
        _heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if _heading_match:
            level = min(6, len(_heading_match.group(1)))
            p = doc.add_heading("", level=level)
            _add_markdown_runs(p, _heading_match.group(2).strip())
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_markdown_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _report_to_pdf_bytes(report_text: str, symbol: str) -> bytes:
    """Convert markdown-like report text to a professional PDF."""
    from fpdf import FPDF
    from datetime import datetime, timezone

    DARK, ACCENT, WHITE = (13, 15, 18), (245, 158, 11), (255, 255, 255)
    GRAY, BODY, BORDER = (140, 140, 140), (50, 50, 50), (220, 220, 220)
    M = 25
    now_str = datetime.now(timezone.utc).strftime("%B %d, %Y at %H:%M UTC")

    def _safe(text: str) -> str:
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        return text.encode("latin-1", errors="ignore").decode("latin-1")

    class ReportPDF(FPDF):
        def header(self):
            if self.page_no() == 1:
                self.set_fill_color(*DARK)
                self.rect(0, 0, 210, 46, style="F")
                self.set_fill_color(*ACCENT)
                self.rect(0, 46, 210, 1.5, style="F")
                self.set_xy(M, 10)
                self.set_text_color(*WHITE)
                self.set_font("Helvetica", style="B", size=20)
                self.cell(w=0, h=8, text="AlphaLens Research Report")
                self.set_xy(M, 22)
                self.set_font("Helvetica", style="B", size=26)
                self.cell(w=0, h=10, text=_safe(symbol))
                self.set_xy(M, 36)
                self.set_text_color(170, 170, 170)
                self.set_font("Helvetica", size=8)
                self.cell(
                    w=0,
                    h=5,
                    text=_safe(
                        f"Generated {now_str}  |  Binance  CoinGecko  DefiLlama  Polymarket"
                    ),
                )
                self.set_y(52)
                self.set_font("Helvetica", style="I", size=7)
                self.set_text_color(*GRAY)
                self.set_x(M)
                self.multi_cell(
                    w=210 - 2 * M,
                    h=4,
                    text=_safe(
                        "For research purposes only. Not financial advice. AI-generated content."
                    ),
                )
                self.ln(4)
            else:
                self.set_y(15)

        def footer(self):
            self.set_y(-14)
            self.set_draw_color(*BORDER)
            self.set_line_width(0.2)
            self.line(M, self.get_y(), 210 - M, self.get_y())
            self.ln(2)
            self.set_font("Helvetica", size=7)
            self.set_text_color(*GRAY)
            cw = 210 - 2 * M
            self.set_x(M)
            self.cell(w=cw / 2, h=4, text="AlphaLens")
            self.cell(w=cw / 2, h=4, text=f"Page {self.page_no()}", align="R")

    pdf = ReportPDF()
    pdf.set_margins(left=M, top=15, right=M)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    cw = 210 - 2 * M

    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            pdf.ln(2)
            continue

        _heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if _heading_match:
            level = len(_heading_match.group(1))
            heading_text = _safe(_heading_match.group(2).strip())
            if level == 1:
                pdf.ln(2)
                pdf.set_font("Helvetica", style="B", size=15)
                pdf.set_text_color(*DARK)
                pdf.set_x(M)
                pdf.multi_cell(w=cw, h=7, text=heading_text)
                pdf.ln(2)
            elif level == 2:
                pdf.ln(4)
                pdf.set_draw_color(*ACCENT)
                pdf.set_line_width(0.4)
                pdf.line(M, pdf.get_y(), M + cw, pdf.get_y())
                pdf.ln(2.5)
                pdf.set_font("Helvetica", style="B", size=13)
                pdf.set_text_color(*DARK)
                pdf.set_x(M)
                pdf.multi_cell(w=cw, h=6.5, text=heading_text)
                pdf.ln(1.5)
            elif level == 3:
                pdf.ln(2)
                pdf.set_font("Helvetica", style="B", size=10.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(M)
                pdf.multi_cell(w=cw, h=5.5, text=heading_text)
                pdf.ln(1)
            else:
                pdf.ln(1.5)
                pdf.set_font("Helvetica", style="B", size=9.8)
                pdf.set_text_color(*DARK)
                pdf.set_x(M)
                pdf.multi_cell(w=cw, h=5, text=heading_text)
                pdf.ln(0.8)
        elif line.startswith(("- ", "* ")):
            pdf.set_font("Helvetica", size=9.5)
            pdf.set_text_color(*BODY)
            pdf.set_x(M + 4)
            pdf.multi_cell(w=cw - 4, h=5, text="- " + _safe(line[2:]))
        else:
            pdf.set_font("Helvetica", size=9.5)
            pdf.set_text_color(*BODY)
            pdf.set_x(M)
            pdf.multi_cell(w=cw, h=5, text=_safe(line))

    return bytes(pdf.output())


# ── Page config ──────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AlphaLens — Crypto Research Agent",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="auto",
)

# ── Theme definitions ─────────────────────────────────────────────────────────────
# Inspired by: Artemis.ai, CoinGlass, Linear, Bloomberg
THEMES = {
    "Studio": {
        "bg": "#0d0f12",
        "sidebar": "#111418",
        "card": "#151820",
        "border": "#1f2937",
        "border_subtle": "#1a1f2b",
        "accent": "#f59e0b",
        "accent_hover": "#fbbf24",
        "accent_text": "#fcd34d",
        "label": "#6b7280",
        "text": "#d1d5db",
        "text_muted": "#9ca3af",
        "chat_bg": "#151820",
        "chat_text": "#d1d5db",
        "code_bg": "#0d0f12",
        "green": "#22c55e",
        "red": "#ef4444",
        "radius": "6px",
    },
    "Studio Light": {
        "bg": "#ffffff",
        "sidebar": "#f9fafb",
        "card": "#ffffff",
        "border": "#e5e7eb",
        "border_subtle": "#f3f4f6",
        "accent": "#d97706",
        "accent_hover": "#b45309",
        "accent_text": "#92400e",
        "label": "#6b7280",
        "text": "#111827",
        "text_muted": "#6b7280",
        "chat_bg": "#f9fafb",
        "chat_text": "#111827",
        "code_bg": "#f3f4f6",
        "green": "#16a34a",
        "red": "#dc2626",
        "radius": "6px",
    },
    "Artemis": {
        "bg": "#09090b",
        "sidebar": "#0c0c0f",
        "card": "#111113",
        "border": "#1e1e22",
        "border_subtle": "#18181b",
        "accent": "#6366f1",
        "accent_hover": "#818cf8",
        "accent_text": "#a5b4fc",
        "label": "#71717a",
        "text": "#e4e4e7",
        "text_muted": "#a1a1aa",
        "chat_bg": "#131316",
        "chat_text": "#e4e4e7",
        "code_bg": "#0c0c0f",
        "green": "#10b981",
        "red": "#ef4444",
        "radius": "8px",
    },
    "Terminal": {
        "bg": "#010409",
        "sidebar": "#0d1117",
        "card": "#0d1117",
        "border": "#21262d",
        "border_subtle": "#161b22",
        "accent": "#22ab94",
        "accent_hover": "#2dd4a8",
        "accent_text": "#6ee7b7",
        "label": "#7d8590",
        "text": "#c9d1d9",
        "text_muted": "#7d8590",
        "chat_bg": "#0d1117",
        "chat_text": "#c9d1d9",
        "code_bg": "#010409",
        "green": "#22ab94",
        "red": "#f23645",
        "radius": "6px",
    },
    "Minimal": {
        "bg": "#000000",
        "sidebar": "#0a0a0a",
        "card": "#0a0a0a",
        "border": "#1a1a1a",
        "border_subtle": "#141414",
        "accent": "#ffffff",
        "accent_hover": "#e5e5e5",
        "accent_text": "#e5e5e5",
        "label": "#666666",
        "text": "#ededed",
        "text_muted": "#888888",
        "chat_bg": "#0a0a0a",
        "chat_text": "#ededed",
        "code_bg": "#050505",
        "green": "#00c853",
        "red": "#ff3d00",
        "radius": "8px",
    },
}

# Prediction Markets — Market Consensus row visuals (badge colors + probability scale)
_CONSENSUS_BADGE_BG = {
    "BTC": "#F7931A",
    "ETH": "#627EEA",
    "SOL": "#9945FF",
    "XRP": "#346AA9",
}

# Signal Scanner — coin badge colors (ticker without USDT)
SCANNER_COIN_BADGE_BG = {
    "BTC": "#F7931A",
    "ETH": "#627EEA",
    "SOL": "#9945FF",
    "XRP": "#346AA9",
    "BNB": "#F3BA2F",
    "ADA": "#0033AD",
    "DOGE": "#C3A634",
}
SCANNER_BADGE_DEFAULT = "#6b7280"
CHATBOT_COIN_COLOR_DOT = {
    "BTC": "🟠",
    "ETH": "🔵",
    "SOL": "🟣",
    "XRP": "🔷",
    "BNB": "🟡",
    "DOGE": "🟨",
    "ADA": "🔵",
    "AVAX": "⚪",
}

_QUICK_RESEARCH_PLACEHOLDER = "Choose a quick prompt…"
_CHATBOT_COIN_OPTIONS = [s.replace("USDT", "") for s in AVAILABLE_SYMBOLS]


def _format_chatbot_coin_option(coin: str) -> str:
    """Render coin option with a colored dot marker."""
    return f"{CHATBOT_COIN_COLOR_DOT.get(coin, '⚪')} {coin}"


def _consensus_prob_bar_color(prob: float) -> str:
    if prob <= 30:
        return "#ef4444"
    if prob <= 60:
        return "#f59e0b"
    return "#22c55e"


def _inject_theme(t: dict) -> None:
    """Inject professional CSS — no gradients, no glow, just clean spacing."""
    r = t["radius"]
    is_light = t["bg"] in ("#ffffff", "#f9fafb", "#fafafa")
    btn_text = (
        "#fff"
        if is_light
        else ("#000" if t["accent"] in ("#ffffff", "#f59e0b") else "#fff")
    )
    st.markdown(
        f"""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        .stApp {{ background: {t["bg"]}; font-family: 'Inter', -apple-system, sans-serif; }}

        /* Sidebar */
        section[data-testid="stSidebar"] {{
            background: {t["sidebar"]};
            border-right: 1px solid {t["border_subtle"]};
        }}
        section[data-testid="stSidebar"] * {{ color: {t["text_muted"]} !important; }}
        section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3, section[data-testid="stSidebar"] strong {{
            color: {t["text"]} !important;
        }}
        section[data-testid="stSidebar"] .stCaption, section[data-testid="stSidebar"] small {{
            color: {t["label"]} !important;
        }}
        /* Single select: paint control + children for contrast. */
        section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] * {{
            color: {t["text"]} !important; background-color: {t["card"]} !important;
        }}
        /*
         * Multiselect: do not set background on every descendant — Base Web tag layout
         * clips the first letter when inner nodes are fully painted (sidebar coin filter).
         */
        section[data-testid="stSidebar"] [data-testid="stMultiSelect"] [data-baseweb="select"] {{
            background-color: {t["card"]} !important;
        }}
        section[data-testid="stSidebar"] [data-testid="stMultiSelect"] [data-baseweb="select"] * {{
            color: {t["text"]} !important;
        }}
        section[data-testid="stSidebar"] button {{
            background: {t["card"]} !important; color: {t["text_muted"]} !important;
            border: 1px solid {t["border"]} !important; border-radius: {r} !important;
            font-size: 0.78rem !important;
        }}
        section[data-testid="stSidebar"] button:hover {{ border-color: {t["accent"]} !important; }}
        section[data-testid="stSidebar"] button *, section[data-testid="stSidebar"] button span {{
            color: {t["text_muted"]} !important; background: transparent !important;
        }}

        /* Metrics */
        [data-testid="stMetric"] {{
            background: {t["card"]}; border: 1px solid {t["border"]};
            border-radius: {r}; padding: 16px 20px 12px;
            min-height: 115px;
        }}
        [data-testid="stMetricLabel"] {{
            font-size: 0.7rem !important; font-weight: 500 !important;
            color: {t["label"]} !important; text-transform: uppercase; letter-spacing: 0.06em;
        }}
        [data-testid="stMetricValue"] {{
            font-size: 1.4rem !important; font-weight: 600 !important;
            color: {t["text"]} !important; font-variant-numeric: tabular-nums;
        }}
        [data-testid="stMetricDelta"] {{ font-size: 0.78rem !important; font-weight: 500 !important; }}
        [data-testid="stMetricDelta"] svg {{ display: none; }}

        /* Metric help icon (native st.metric help=) — subtle, matches prior “i” intent */
        [data-testid="stMetric"] [data-testid="stMetricLabel"] button {{
            color: #6b7280 !important;
        }}
        [data-testid="stMetric"] [data-testid="stMetricLabel"] button:hover {{
            color: {t["text"]} !important;
        }}

        /* Tabs — clean underline style */
        .stTabs [data-baseweb="tab-list"] {{
            gap: 0; background: transparent; border-radius: 0;
            padding: 0; border: none; border-bottom: 1px solid {t["border"]};
        }}
        .stTabs [data-baseweb="tab"] {{
            border-radius: 0; padding: 10px 20px; font-weight: 500;
            font-size: 0.84rem; color: {t["text_muted"]}; letter-spacing: 0;
        }}
        .stTabs [aria-selected="true"] {{
            background: transparent !important; color: {t["text"]} !important;
            border-bottom: 2px solid {t["accent"]} !important;
        }}
        .stTabs [data-baseweb="tab-highlight"] {{ display: none; }}
        .stTabs [data-baseweb="tab-border"] {{ display: none; }}

        /* Expanders */
        .streamlit-expanderHeader {{
            background: {t["card"]} !important; border-radius: {r} !important;
            border: 1px solid {t["border"]} !important; font-weight: 500 !important;
        }}

        /* Dataframes */
        [data-testid="stDataFrame"] {{
            border-radius: {r}; overflow: hidden; border: 1px solid {t["border"]};
        }}

        /* Dividers */
        hr {{ border-color: {t["border_subtle"]} !important; }}

        /* Chat */
        [data-testid="stChatMessage"] {{
            background: {t["chat_bg"]}; border: 1px solid {t["border"]};
            border-radius: {r}; padding: 16px 20px; margin-bottom: 8px;
        }}
        [data-testid="stChatMessage"] *:not(svg):not(path) {{ color: {t["chat_text"]} !important; }}
        [data-testid="stChatMessage"] h1, [data-testid="stChatMessage"] h2,
        [data-testid="stChatMessage"] h3, [data-testid="stChatMessage"] strong {{
            color: {t["text"]} !important; font-weight: 600 !important;
        }}
        [data-testid="stChatMessage"] code {{
            color: {t["accent_text"]} !important; background: {t["code_bg"]} !important;
            border-radius: 4px; padding: 2px 5px; font-size: 0.85em;
        }}
        [data-testid="stChatMessage"] pre {{
            background: {t["code_bg"]} !important; border-radius: {r};
            padding: 12px 16px; border: 1px solid {t["border_subtle"]};
        }}
        [data-testid="stChatMessage"] pre code {{ color: {t["chat_text"]} !important; background: transparent !important; padding: 0; }}
        [data-testid="stChatMessage"] blockquote {{ border-left: 2px solid {t["border"]}; padding-left: 12px; margin-left: 0; }}
        [data-testid="stChatMessage"] a {{ color: {t["accent_text"]} !important; }}
        .stChatInputContainer {{
            border-color: {t["border"]} !important; border-radius: {r} !important;
            background: {t["card"]} !important;
        }}
        .stChatInputContainer textarea {{ color: {t["text"]} !important; }}

        /* Hide Streamlit default UI chrome */
        #MainMenu {{ visibility: hidden; }}
        header[data-testid="stHeader"] button[kind="header"] {{ display: none; }}
        [data-testid="stDecoration"] {{ display: none; }}
        footer {{ display: none !important; }}
        .stDeployButton {{ display: none !important; }}

        /* Scrollbar */
        [data-testid="stVerticalBlockBorderWrapper"] {{ scrollbar-width: thin; scrollbar-color: {t["border"]} transparent; }}
        [data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar {{ width: 4px; }}
        [data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar-thumb {{ background: {t["border"]}; border-radius: 4px; }}

        /* Buttons */
        .stButton > button[kind="primary"] {{
            background: {t["accent"]}; color: {btn_text} !important;
            border: none; border-radius: {r}; font-weight: 600;
            transition: background 0.15s ease, transform 0.1s ease;
        }}
        .stButton > button[kind="primary"]:hover {{ background: {t["accent_hover"]}; transform: translateY(-1px); }}
        .stButton > button[kind="primary"]:active {{ transform: translateY(0); }}
        .stButton > button {{
            transition: border-color 0.15s ease, background 0.15s ease;
        }}

        /* Cards / Metrics — smooth hover */
        [data-testid="stMetric"] {{
            transition: border-color 0.15s ease, box-shadow 0.15s ease;
        }}
        [data-testid="stMetric"]:hover {{
            border-color: {t["accent"]}40;
            box-shadow: 0 0 0 1px {t["accent"]}20;
        }}

        /* Download buttons — compact style */
        .stDownloadButton > button {{
            background: {t["card"]} !important;
            border: 1px solid {t["border"]} !important;
            border-radius: {r} !important;
            color: {t["text_muted"]} !important;
            font-size: 0.78rem !important;
            font-weight: 500 !important;
            transition: border-color 0.15s ease, color 0.15s ease;
        }}
        .stDownloadButton > button:hover {{
            border-color: {t["accent"]} !important;
            color: {t["text"]} !important;
        }}

        /* Tab transition */
        .stTabs [data-baseweb="tab"] {{
            transition: color 0.15s ease;
        }}

        /* Chat input focus glow */
        .stChatInputContainer:focus-within {{
            border-color: {t["accent"]}80 !important;
            box-shadow: 0 0 0 1px {t["accent"]}30;
        }}

        /* Status widget — cleaner */
        [data-testid="stStatusWidget"] {{
            background: {t["card"]} !important;
            border: 1px solid {t["border"]} !important;
            border-radius: {r} !important;
        }}

        /* Selectbox hover */
        [data-baseweb="select"] {{
            transition: border-color 0.15s ease;
        }}

        /* Anchor links on headings — hide by default */
        [data-testid="stHeadingWithActionElements"] a {{
            opacity: 0;
            transition: opacity 0.15s ease;
        }}
        [data-testid="stHeadingWithActionElements"]:hover a {{
            opacity: 0.5;
        }}

        /* Status badge (dashboard live strip — keep clear of metrics below) */
        @keyframes alphalens-live-pulse {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.4; }}
        }}
        .status-bar {{
            display: inline-flex; align-items: center; gap: 8px;
            background: {t["card"]}; border: 1px solid {t["border"]};
            border-radius: 20px; padding: 5px 14px;
            margin-bottom: 1rem;
            font-size: 0.75rem; color: {t["text_muted"]}; font-weight: 500;
        }}
        .status-bar .dot {{
            width: 6px; height: 6px; border-radius: 50%;
            background: {t["green"]}; display: inline-block;
            animation: alphalens-live-pulse 2s ease-in-out infinite;
        }}
        .status-bar .dot.off {{ background: {t["red"]}; animation: none; }}

        /* Market Consensus card — row layout + live dot (Prediction Markets tab) */
        @keyframes alphalens-consensus-pulse {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.42; }}
        }}
        .alphalens-consensus-live-dot {{
            width: 6px; height: 6px; border-radius: 50%;
            background: {t["green"]}; flex-shrink: 0;
            animation: alphalens-consensus-pulse 1.8s ease-in-out infinite;
        }}
        .alphalens-consensus-row {{
            display: flex; align-items: center; gap: 14px;
            padding: 12px 16px;
            border-bottom: {"1px solid rgba(0,0,0,0.08)" if is_light else "1px solid rgba(55, 65, 81, 0.5)"};
            transition: background 0.12s ease;
        }}
        .alphalens-consensus-row:last-child {{ border-bottom: none; }}
        .alphalens-consensus-row:hover {{
            background: {"rgba(0,0,0,0.04)" if is_light else "rgba(255,255,255,0.05)"};
        }}

        /* Containers with border — prediction market cards */
        [data-testid="stVerticalBlockBorderWrapper"]:has(> div > [data-testid="stVerticalBlock"]) {{
            border-color: {t["border"]} !important;
            border-radius: {r} !important;
        }}

        /* Signal Scanner table */
        .alphalens-scanner-wrap {{
            background: {t["card"]};
            border: 1px solid {t["border"]};
            border-radius: {r};
            overflow: hidden;
            margin-bottom: 1rem;
        }}
        .alphalens-scanner-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
        }}
        .alphalens-scanner-table thead th {{
            text-align: left;
            padding: 12px 16px;
            font-size: 12px;
            font-weight: 600;
            letter-spacing: 0.06em;
            text-transform: uppercase;
            color: {t["label"]};
            border-bottom: 1px solid {t["border"]};
            background: {t["card"]};
        }}
        .alphalens-scanner-row td {{
            padding: 14px 18px;
            vertical-align: top;
            border-bottom: 1px solid {t["border_subtle"]};
        }}
        .alphalens-scanner-cell-main {{
            font-weight: 600;
            font-size: 0.9375rem;
            line-height: 1.3;
        }}
        .alphalens-scanner-cell-sub {{
            font-size: 0.8125rem;
            font-weight: 500;
            line-height: 1.45;
            margin-top: 6px;
            letter-spacing: 0.01em;
        }}
        .alphalens-scanner-trend {{
            font-weight: 600;
            font-size: 0.9375rem;
        }}
        .alphalens-scanner-row:nth-child(even) td {{
            background: {"rgba(0,0,0,0.02)" if is_light else "rgba(255,255,255,0.02)"};
        }}
        .alphalens-scanner-row:hover td {{
            background: {"rgba(0,0,0,0.04)" if is_light else "rgba(255,255,255,0.05)"} !important;
        }}
        .alphalens-scanner-pill {{
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 5px 12px;
            border-radius: 9999px;
            font-size: 12px;
            font-weight: 600;
            border: 1px solid {t["border"]};
        }}
        .alphalens-scanner-summary {{
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin: 1rem 0 1.25rem;
        }}

        /* News feed cards */
        .alphalens-news-card {{
            padding: 16px 18px;
            margin-bottom: 10px;
            border-radius: {r};
            border: 1px solid {t["border"]};
            background: {t["card"]};
            transition: border-color 0.15s ease, transform 0.1s ease;
        }}
        .alphalens-news-card:hover {{
            border-color: {t["accent"]}40;
            transform: translateY(-1px);
        }}
        .alphalens-news-card .news-source {{
            font-size: 0.72rem; font-weight: 600;
            color: {t["accent"]}; letter-spacing: 0.04em;
            text-transform: uppercase;
        }}
        .alphalens-news-card .news-time {{
            font-size: 0.7rem; color: {t["text_muted"]}; opacity: 0.6;
        }}
        .alphalens-news-card .news-title {{
            font-size: 0.92rem; font-weight: 600; color: {t["text"]};
            line-height: 1.4; margin-bottom: 6px;
            transition: color 0.15s ease;
        }}
        .alphalens-news-card .news-title:hover {{
            color: {t["accent_text"]};
        }}
        .alphalens-news-card .news-body {{
            font-size: 0.78rem; color: {t["text_muted"]}; opacity: 0.7;
            line-height: 1.5;
        }}
        .alphalens-news-card .news-footer {{
            display: flex; justify-content: flex-end; margin-top: 6px;
            padding-top: 6px; border-top: 1px solid {t["border_subtle"]};
        }}
        .alphalens-news-card .news-watermark {{
            font-size: 0.68rem; color: {t["text_muted"]}; opacity: 0.35;
            font-weight: 500; letter-spacing: 0.03em;
        }}

        /* ── Mobile responsive ──────────────────────────────────────── */

        /* Tabs: horizontal scroll on small screens */
        @media (max-width: 768px) {{
            .stTabs [data-baseweb="tab-list"] {{
                overflow-x: auto;
                -webkit-overflow-scrolling: touch;
                scrollbar-width: none;
                flex-wrap: nowrap;
            }}
            .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar {{ display: none; }}
            .stTabs [data-baseweb="tab"] {{
                padding: 8px 12px;
                font-size: 0.75rem;
                white-space: nowrap;
                flex-shrink: 0;
            }}

            /* Stack all Streamlit column layouts vertically */
            [data-testid="stHorizontalBlock"] {{
                flex-direction: column !important;
                gap: 8px !important;
            }}
            [data-testid="stColumn"] {{
                width: 100% !important;
                flex: 1 1 100% !important;
                min-width: 0 !important;
            }}

            /* Metric cards: compact on mobile */
            [data-testid="stMetric"] {{
                padding: 10px 14px 8px;
                min-height: 80px;
            }}
            [data-testid="stMetricValue"] {{
                font-size: 1.15rem !important;
            }}
            [data-testid="stMetricLabel"] {{
                font-size: 0.65rem !important;
            }}

            /* Chat messages: tighter padding */
            [data-testid="stChatMessage"] {{
                padding: 12px 14px;
            }}

            /* Scanner table: horizontal scroll + smaller cells */
            .alphalens-scanner-wrap {{
                overflow-x: auto;
                -webkit-overflow-scrolling: touch;
            }}
            .alphalens-scanner-table {{
                min-width: 580px;
                font-size: 13px;
            }}
            .alphalens-scanner-table thead th {{
                padding: 10px 12px;
                font-size: 11px;
            }}
            .alphalens-scanner-row td {{
                padding: 10px 12px;
            }}
            .alphalens-scanner-cell-main {{
                font-size: 0.875rem;
            }}
            .alphalens-scanner-cell-sub {{
                font-size: 0.75rem;
                margin-top: 5px;
            }}

            /* Consensus card: stack probability bar below text */
            .alphalens-consensus-row {{
                flex-wrap: wrap;
                gap: 8px;
                padding: 10px 12px;
            }}
            .alphalens-consensus-row > div:last-child {{
                width: 100% !important;
                max-width: 100% !important;
            }}

            /* Status bar: wrap on narrow screens */
            .status-bar {{
                flex-wrap: wrap;
                font-size: 0.68rem;
                padding: 4px 10px;
            }}

            /* Main content area: reduce side padding */
            .stMainBlockContainer {{
                padding-left: 1rem !important;
                padding-right: 1rem !important;
            }}

            /* Plotly charts: reduce margin */
            .js-plotly-plot {{
                margin-left: -8px;
                margin-right: -8px;
            }}
        }}

        /* Small phones */
        @media (max-width: 480px) {{
            [data-testid="stMetric"] {{
                padding: 8px 10px 6px;
                min-height: 70px;
            }}
            [data-testid="stMetricValue"] {{
                font-size: 1rem !important;
            }}
            .stTabs [data-baseweb="tab"] {{
                padding: 6px 10px;
                font-size: 0.7rem;
            }}
            [data-testid="stChatMessage"] {{
                padding: 10px 12px;
            }}
            .alphalens-consensus-row {{
                padding: 8px 10px;
            }}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


# ── One-time DB init ─────────────────────────────────────────────────────────────
@st.cache_resource
def setup_db():
    init_db()


setup_db()

# ── Session state defaults ───────────────────────────────────────────────────────
# Only pre-init keys that are NOT owned by a widget (widgets set their own defaults).
if "history_loaded" not in st.session_state:
    st.session_state["history_loaded"] = set()
if "chat_messages" not in st.session_state:
    st.session_state["chat_messages"] = []
if "conversations" not in st.session_state:
    st.session_state["conversations"] = []  # [{id, title, messages}]
if "active_conv_id" not in st.session_state:
    st.session_state["active_conv_id"] = None
if "active_tab" not in st.session_state:
    st.session_state["active_tab"] = "Dashboard"
if "loaded_history" not in st.session_state:
    st.session_state["loaded_history"] = set()

SCANNER_DATA_LIMIT = 60
SCANNER_CACHE_TTL = 30


def _scanner_fetch_symbol_worker(
    symbol: str,
    interval: str,
    limit: int,
    loaded_snapshot: frozenset,
):
    """
    Fetch klines + indicators for one symbol. Runs in a thread pool — do not use
    st.session_state here; backfill keys are returned for the main thread to record.
    """
    from src.database import get_klines
    from src.history import fetch_historical_klines

    key = f"{symbol}_{interval}"
    try:
        raw = list(get_klines(symbol, interval=interval, limit=limit)) or []
        backfill_key = None
        if len(raw) < 20 and key not in loaded_snapshot:
            fetch_historical_klines(symbol, interval, limit=limit)
            raw = list(get_klines(symbol, interval=interval, limit=limit)) or []
            backfill_key = key

        if len(raw) < 20:
            return symbol, None, backfill_key

        df = pd.DataFrame(
            raw,
            columns=["open_time", "open", "high", "low", "close", "volume"],
        )
        df = add_indicators(df)
        if df.empty or len(df) < 2:
            return symbol, None, backfill_key
        return symbol, df, backfill_key
    except Exception:
        return symbol, None, None


def _scanner_get_fresh_cache(interval: str, force_refresh: bool):
    """Return scanner_cache dict if TTL-valid and interval matches; else None."""
    if force_refresh:
        return None
    cache = st.session_state.get("scanner_cache") or {}
    ts = float(st.session_state.get("scanner_cache_ts", 0) or 0)
    if cache.get("interval") != interval:
        return None
    if time.time() - ts >= SCANNER_CACHE_TTL:
        return None
    if not cache.get("rows"):
        return None
    return cache


def _scanner_set_display_cache(interval: str, rows_out: list, ts_str: str) -> None:
    st.session_state["scanner_cache"] = {
        "interval": interval,
        "rows": rows_out,
        "ts_str": ts_str,
    }
    st.session_state["scanner_cache_ts"] = time.time()


def _scanner_signal_counts(
    rsi,
    macd,
    sig,
    bb_pct,
    sma20,
    sma50,
) -> tuple[int, int, int]:
    """Bullish / bearish / neutral counts (0–1 each) across RSI, MACD, BB, trend."""
    bull = bear = neutral = 0

    if rsi is None or pd.isna(rsi):
        neutral += 1
    elif float(rsi) < 30:
        bull += 1
    elif float(rsi) > 70:
        bear += 1
    else:
        neutral += 1

    if macd is None or sig is None or pd.isna(macd) or pd.isna(sig):
        neutral += 1
    elif float(macd) > float(sig):
        bull += 1
    elif float(macd) < float(sig):
        bear += 1
    else:
        neutral += 1

    if bb_pct is None or pd.isna(bb_pct):
        neutral += 1
    elif float(bb_pct) < 20:
        bull += 1
    elif float(bb_pct) > 80:
        bear += 1
    else:
        neutral += 1

    if sma20 is None or sma50 is None or pd.isna(sma20) or pd.isna(sma50):
        neutral += 1
    else:
        s20, s50 = float(sma20), float(sma50)
        if s50 == 0:
            neutral += 1
        elif abs(s20 - s50) / abs(s50) < 0.001:
            neutral += 1
        elif s20 > s50:
            bull += 1
        else:
            bear += 1

    return bull, bear, neutral


def _scanner_overall_label(bull: int, bear: int) -> str:
    if bull >= 3:
        return "🟢 Strong Buy"
    if bear >= 3:
        return "🔴 Strong Sell"
    if bull == 2 and bear == 0:
        return "🟡 Buy"
    if bear == 2 and bull == 0:
        return "🟡 Sell"
    return "⚪ Neutral"


def _archive_current_chat() -> None:
    """Save or update the current chat into the conversations list."""
    import time as _time

    msgs = st.session_state.get("chat_messages", [])
    if not msgs:
        return
    user_msgs = [m for m in msgs if m["role"] == "user"]
    title = (user_msgs[0]["content"][:42] + "…") if user_msgs else "Untitled"
    conv_id = st.session_state.get("active_conv_id")
    convs = st.session_state["conversations"]
    if conv_id:
        for c in convs:
            if c["id"] == conv_id:
                c["messages"] = [dict(m) for m in msgs]
                c["title"] = title
                return
    new_id = str(int(_time.time() * 1000))
    convs.append({"id": new_id, "title": title, "messages": [dict(m) for m in msgs]})
    st.session_state["active_conv_id"] = new_id
    # Keep only the 10 most recent conversations to bound memory
    if len(convs) > 10:
        st.session_state["conversations"] = convs[-10:]


def _run_selected_quick_research() -> None:
    """Queue selected quick prompt with selected coin for chatbot execution."""
    from src.prompts.quick_prompts import QUICK_PROMPTS as _qp

    sel = st.session_state.get("quick_research_dropdown")
    if not sel or sel == _QUICK_RESEARCH_PLACEHOLDER:
        return
    _sym = str(st.session_state.get("quick_research_coin") or "").strip().upper()
    if _sym not in _CHATBOT_COIN_OPTIONS:
        _pa = (st.session_state.get("dash_assets") or ["BTCUSDT"])[0]
        _sym = str(_pa).replace("USDT", "")
    for _l, _e, _tmpl in _qp:
        if f"{_e} {_l}" == sel:
            st.session_state["_chat_user_input"] = _tmpl.format(symbol=_sym)
            break
    # Reset after rerun, before selectbox is instantiated again.
    st.session_state["_reset_quick_research_dropdown"] = True


# ── Sidebar (branding + theme) ─────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📈 AlphaLens")
    st.caption("AI-Powered Crypto Research Agent")
    st.divider()

    theme_name: str = st.selectbox(
        "Theme",
        list(THEMES.keys()),
        index=0,
        key="theme_select",
    )

_active_theme = THEMES[theme_name]
_inject_theme(_active_theme)
_is_light_theme = _active_theme["bg"] in ("#ffffff", "#f9fafb", "#fafafa")
_plotly_template = "plotly_white" if _is_light_theme else "plotly_dark"
_grid_color = "#e5e7eb" if _is_light_theme else "#2a2a2a"

# ════════════════════════════════════════════════════════════════════════════════
# TABS (main area — horizontal bar unchanged)
# ════════════════════════════════════════════════════════════════════════════════
# Streamlit 1.50: st.tabs() only accepts a list of tab titles (no extra kwargs).

tab_dashboard, tab_research, tab_signal_scanner, tab_prediction, tab_news = st.tabs(
    [
        "📊 Dashboard",
        "💬 Chatbot",
        "🔍 Signal Scanner",
        "🎯 Prediction Markets",
        "📰 News Feed",
    ]
)

# Tab tracking via st.tabs context (tab.open not available in Streamlit 1.50)
if "active_tab" not in st.session_state:
    st.session_state["active_tab"] = "Dashboard"

# ── Sidebar (dashboard controls + footer) ──────────────────────────────────────
# NOTE: st.tabs() is layout-only — every `with tab_X:` block runs on every rerun,
# so we can't reliably know which tab the user is viewing. Earlier versions tried
# to key sidebar contents off st.session_state["active_tab"], but because each tab
# body overwrote that value (with the last one — News Feed — winning), the sidebar
# would go blank after the first rerun. Dashboard controls now always render here;
# per-tab controls (chatbot conversations, scanner refresh, market filters) live
# inside their own tab bodies where they only render when that tab is active.
with st.sidebar:
    st.divider()
    st.multiselect(
        "Assets (up to 4)",
        AVAILABLE_SYMBOLS,
        default=["BTCUSDT"],
        max_selections=4,
        key="dash_assets",
    )
    _di = st.session_state.get("dash_interval", "1m")
    _di_ix = _DASH_INTERVAL_OPTS.index(_di) if _di in _DASH_INTERVAL_OPTS else 0
    st.selectbox(
        "Candle interval",
        _DASH_INTERVAL_OPTS,
        index=_di_ix,
        key="dash_interval",
    )
    st.slider(
        "Candles to display",
        min_value=20,
        max_value=500,
        value=int(st.session_state.get("dash_lookback", 100)),
        step=10,
        key="dash_lookback",
    )
    st.divider()
    st.toggle("Auto-refresh (5 s)", key="auto_refresh")

    st.divider()
    st.caption("Data: Binance · CoinGecko · DefiLlama · Polymarket")
    st.caption("AI: GPT-4o / Gemini")

symbols = list(st.session_state.get("dash_assets") or ["BTCUSDT"])
if not symbols:
    symbols = ["BTCUSDT"]
interval = str(st.session_state.get("dash_interval", "1m"))
lookback = int(st.session_state.get("dash_lookback", 100))
auto_refresh = bool(st.session_state.get("auto_refresh", True))

primary = symbols[0]

# ── Historical data: fetch once per (symbol, interval), parallel ──────────────
_syms_to_fetch = [
    sym
    for sym in symbols
    if f"hist_{sym}_{interval}" not in st.session_state["history_loaded"]
]
if _syms_to_fetch:
    with st.spinner(f"Loading historical candles for {', '.join(_syms_to_fetch)}…"):
        with ThreadPoolExecutor(max_workers=4) as _ex:
            results = list(
                _ex.map(
                    lambda s: (s, fetch_historical_klines(s, interval, limit=500)),
                    _syms_to_fetch,
                )
            )
    for sym, count in results:
        if count > 0:
            st.session_state["history_loaded"].add(f"hist_{sym}_{interval}")

# ── Kline WebSocket streams ──────────────────────────────────────────────────────
active_ks = {f"ks_{sym}_{interval}" for sym in symbols}

for k in list(st.session_state.keys()):
    if k.startswith("ks_") and k not in active_ks:
        v = st.session_state[k]
        if isinstance(v, BinanceKlineStream):
            v.stop()
            del st.session_state[k]

for sym in symbols:
    ks_key = f"ks_{sym}_{interval}"
    if ks_key not in st.session_state:
        s = BinanceKlineStream(sym, interval)
        s.start()
        st.session_state[ks_key] = s

# ── Order book streams: all selected symbols ─────────────────────────────────────
active_obs = {f"ob_{sym}" for sym in symbols}

for k in list(st.session_state.keys()):
    if k.startswith("ob_") and k not in active_obs:
        v = st.session_state[k]
        if isinstance(v, BinanceOrderBookStream):
            v.stop()
            del st.session_state[k]

for sym in symbols:
    ob_key = f"ob_{sym}"
    if ob_key not in st.session_state:
        obs = BinanceOrderBookStream(sym)
        obs.start()
        st.session_state[ob_key] = obs

# ── Header ───────────────────────────────────────────────────────────────────────
primary_stream = st.session_state.get(f"ks_{primary}_{interval}")
is_live = primary_stream and primary_stream.is_running

# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DASHBOARD (charts + order book)
# Fragment re-renders every 5s independently — chat and deep dive are untouched.
# ════════════════════════════════════════════════════════════════════════════════

with tab_dashboard:
    if True:  # tab_dashboard
        st.session_state["active_tab"] = "Dashboard"
        _asset_title = " · ".join(symbols) if len(symbols) > 1 else primary
        st.markdown(f"### {_asset_title}")
        _dot_cls = "" if is_live else " off"
        _status_text = "Live" if is_live else "Disconnected"
        st.markdown(
            f'<div class="status-bar">'
            f'<span class="dot{_dot_cls}"></span> {_status_text}'
            f" &nbsp;·&nbsp; {interval} candles"
            f" &nbsp;·&nbsp; {lookback} loaded"
            f" &nbsp;·&nbsp; {len(symbols)} stream(s)"
            f"</div>",
            unsafe_allow_html=True,
        )

        @st.fragment(run_every=5 if auto_refresh else None)
        def _dashboard():
            from src.database import get_klines
            from src.indicators import add_indicators

            dfs: dict[str, pd.DataFrame] = {}
            for _s in symbols:
                _k = get_klines(_s, interval=interval, limit=lookback)
                if not _k:
                    fetch_historical_klines(_s, interval, limit=500)
                    _k = get_klines(_s, interval=interval, limit=lookback)
                if _k:
                    _d = pd.DataFrame(
                        _k,
                        columns=["open_time", "open", "high", "low", "close", "volume"],
                    )
                    _d["datetime"] = pd.to_datetime(
                        _d["open_time"], unit="ms", utc=True
                    )
                    _d = add_indicators(_d)
                    dfs[_s] = _d

            if len(symbols) == 1:
                df = dfs.get(primary)
                if df is None or df.empty:
                    st.warning("Could not load candle data. Retrying on next refresh…")
                    return

                latest = df.iloc[-1]
                earliest = df.iloc[0]
                pct = ((latest["close"] - earliest["close"]) / earliest["close"]) * 100
                dollar_chg = latest["close"] - earliest["close"]

                # Row 1 — price metrics
                c1, c2, c3, c4 = st.columns(4)
                render_metric_with_tooltip(
                    c1,
                    "Price (USDT)",
                    f"${latest['close']:,.2f}",
                    f"{dollar_chg:+,.2f} ({pct:+.2f}%)",
                )
                render_metric_with_tooltip(
                    c2,
                    "Period High",
                    f"${df['high'].max():,.2f}",
                )
                render_metric_with_tooltip(
                    c3,
                    "Period Low",
                    f"${df['low'].min():,.2f}",
                )
                render_metric_with_tooltip(
                    c4,
                    "Volume",
                    f"{df['volume'].sum():,.0f}",
                )

                # Row 2 — technical indicators
                rsi_val = latest.get("rsi")
                macd_val = latest.get("macd")
                sig_val = latest.get("macd_signal")
                bb_up = latest.get("bb_upper")
                bb_lo = latest.get("bb_lower")

                t1, t2, t3 = st.columns(3)
                render_metric_with_tooltip(
                    t1,
                    "RSI (14)",
                    f"{rsi_val:.1f}" if rsi_val and not pd.isna(rsi_val) else "—",
                    "overbought"
                    if (rsi_val and rsi_val > 70)
                    else "oversold"
                    if (rsi_val and rsi_val < 30)
                    else "neutral",
                )
                render_metric_with_tooltip(
                    t2,
                    "MACD",
                    f"{macd_val:.4f}" if macd_val and not pd.isna(macd_val) else "—",
                    "▲ bullish"
                    if (macd_val and sig_val and macd_val > sig_val)
                    else "▼ bearish",
                )
                if bb_up and bb_lo and not pd.isna(bb_up):
                    bw = bb_up - bb_lo
                    bpos = (latest["close"] - bb_lo) / bw * 100 if bw > 0 else 50
                    render_metric_with_tooltip(
                        t3,
                        "BB Position",
                        f"{bpos:.0f}%",
                        "near top"
                        if bpos > 80
                        else "near bottom"
                        if bpos < 20
                        else "mid-band",
                    )
                else:
                    render_metric_with_tooltip(t3, "BB Position", "—")

                fig = go.Figure()
                fig.add_trace(
                    go.Scatter(
                        x=df["datetime"],
                        y=df["bb_upper"],
                        mode="lines",
                        line=dict(color="rgba(255,200,0,0.3)", width=1),
                        name="BB Upper",
                    )
                )
                fig.add_trace(
                    go.Scatter(
                        x=df["datetime"],
                        y=df["bb_lower"],
                        mode="lines",
                        line=dict(color="rgba(255,200,0,0.3)", width=1),
                        fill="tonexty",
                        fillcolor="rgba(255,200,0,0.05)",
                        name="BB Lower",
                    )
                )
                fig.add_trace(
                    go.Scatter(
                        x=df["datetime"],
                        y=df["bb_mid"],
                        mode="lines",
                        line=dict(color="rgba(255,200,0,0.5)", width=1, dash="dot"),
                        name="BB Mid",
                        showlegend=False,
                    )
                )
                fig.add_trace(
                    go.Candlestick(
                        x=df["datetime"],
                        open=df["open"],
                        high=df["high"],
                        low=df["low"],
                        close=df["close"],
                        name="Price",
                        increasing_line_color="#00e676",
                        decreasing_line_color="#ff1744",
                        increasing_fillcolor="#00e676",
                        decreasing_fillcolor="#ff1744",
                    )
                )
                fig.add_trace(
                    go.Scatter(
                        x=df["datetime"],
                        y=df["sma_20"],
                        mode="lines",
                        line=dict(color="#7c83fd", width=1.5),
                        name="SMA 20",
                    )
                )
                fig.add_trace(
                    go.Scatter(
                        x=df["datetime"],
                        y=df["sma_50"],
                        mode="lines",
                        line=dict(color="#ff9800", width=1.5),
                        name="SMA 50",
                    )
                )
                fig.add_trace(
                    go.Bar(
                        x=df["datetime"],
                        y=df["volume"],
                        name="Volume",
                        yaxis="y2",
                        marker_color=[
                            "rgba(0,230,118,0.2)" if c >= o else "rgba(255,23,68,0.2)"
                            for o, c in zip(df["open"], df["close"])
                        ],
                    )
                )

                fig.update_layout(
                    template=_plotly_template,
                    height=540,
                    margin=dict(t=78, l=60, r=56, b=40),
                    title=dict(
                        text=f"{primary} · {interval}  |  SMA 20  SMA 50  BB",
                        x=0.01,
                        xanchor="left",
                        pad=dict(b=10),
                        font=dict(size=13, color="#aaaaaa"),
                    ),
                    xaxis=dict(showgrid=True, gridcolor=_grid_color),
                    yaxis=dict(
                        title="Price (USDT)", showgrid=True, gridcolor=_grid_color
                    ),
                    yaxis2=dict(
                        title="Volume", overlaying="y", side="right", showgrid=False
                    ),
                    xaxis_rangeslider_visible=False,
                    legend=dict(
                        orientation="h",
                        yanchor="bottom",
                        y=1.02,
                        xanchor="right",
                        x=0.99,
                        font=dict(size=11),
                        bgcolor="rgba(0,0,0,0)",
                    ),
                    plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)",
                )
                st.plotly_chart(fig, width="stretch")

                with st.expander("📊 Technical Indicators — RSI & MACD", expanded=True):
                    fig_ind = make_subplots(
                        rows=2,
                        cols=1,
                        shared_xaxes=True,
                        row_heights=[0.5, 0.5],
                        vertical_spacing=0.08,
                        subplot_titles=("RSI (14)", "MACD (12, 26, 9)"),
                    )
                    fig_ind.add_trace(
                        go.Scatter(
                            x=df["datetime"],
                            y=df["rsi"],
                            mode="lines",
                            line=dict(color="#7c83fd", width=2),
                            name="RSI",
                        ),
                        row=1,
                        col=1,
                    )
                    fig_ind.add_hline(
                        y=70,
                        line_dash="dash",
                        line_color="rgba(255,23,68,0.5)",
                        annotation_text="Overbought 70",
                        row=1,
                        col=1,
                    )
                    fig_ind.add_hline(
                        y=30,
                        line_dash="dash",
                        line_color="rgba(0,230,118,0.5)",
                        annotation_text="Oversold 30",
                        row=1,
                        col=1,
                    )
                    fig_ind.add_hline(
                        y=50,
                        line_dash="dot",
                        line_color="rgba(255,255,255,0.15)",
                        row=1,
                        col=1,
                    )
                    hist_colors = [
                        "rgba(0,230,118,0.6)" if v >= 0 else "rgba(255,23,68,0.6)"
                        for v in df["macd_hist"].fillna(0)
                    ]
                    fig_ind.add_trace(
                        go.Bar(
                            x=df["datetime"],
                            y=df["macd_hist"],
                            name="Histogram",
                            marker_color=hist_colors,
                        ),
                        row=2,
                        col=1,
                    )
                    fig_ind.add_trace(
                        go.Scatter(
                            x=df["datetime"],
                            y=df["macd"],
                            mode="lines",
                            line=dict(color="#7c83fd", width=2),
                            name="MACD",
                        ),
                        row=2,
                        col=1,
                    )
                    fig_ind.add_trace(
                        go.Scatter(
                            x=df["datetime"],
                            y=df["macd_signal"],
                            mode="lines",
                            line=dict(color="#ff9800", width=1.5, dash="dot"),
                            name="Signal",
                        ),
                        row=2,
                        col=1,
                    )
                    fig_ind.add_hline(
                        y=0, line_color="rgba(255,255,255,0.15)", row=2, col=1
                    )
                    fig_ind.update_layout(
                        template=_plotly_template,
                        height=420,
                        margin=dict(t=56, l=60, r=40, b=64),
                        plot_bgcolor="rgba(0,0,0,0)",
                        paper_bgcolor="rgba(0,0,0,0)",
                        legend=dict(
                            orientation="h",
                            yanchor="top",
                            y=-0.12,
                            xanchor="center",
                            x=0.5,
                            font=dict(size=11),
                            bgcolor="rgba(0,0,0,0)",
                        ),
                        showlegend=True,
                    )
                    fig_ind.update_yaxes(showgrid=True, gridcolor=_grid_color)
                    fig_ind.update_xaxes(showgrid=True, gridcolor=_grid_color)
                    st.plotly_chart(fig_ind, width="stretch")

            else:
                metric_cols = st.columns(len(symbols))
                for i, sym in enumerate(symbols):
                    df = dfs.get(sym)
                    if df is not None and not df.empty:
                        latest = df.iloc[-1]
                        earliest = df.iloc[0]
                        pct = (
                            (latest["close"] - earliest["close"]) / earliest["close"]
                        ) * 100
                        metric_cols[i].metric(
                            sym, f"${latest['close']:,.4f}", f"{pct:+.2f}%"
                        )
                    else:
                        metric_cols[i].metric(sym, "—", "loading…")

                fig_comp = go.Figure()
                for i, sym in enumerate(symbols):
                    df = dfs.get(sym)
                    if df is not None and not df.empty:
                        base = df["close"].iloc[0]
                        normalized = (df["close"] / base - 1) * 100
                        fig_comp.add_trace(
                            go.Scatter(
                                x=df["datetime"],
                                y=normalized,
                                mode="lines",
                                name=sym,
                                line=dict(
                                    color=CHART_COLORS[i % len(CHART_COLORS)], width=2
                                ),
                            )
                        )
                fig_comp.add_hline(
                    y=0, line_dash="dash", line_color="rgba(255,255,255,0.15)"
                )
                fig_comp.update_layout(
                    template=_plotly_template,
                    height=360,
                    margin=dict(t=78, l=60, r=56, b=40),
                    title=dict(
                        text="Relative % Change (normalized to period start = 0%)",
                        x=0.01,
                        xanchor="left",
                        pad=dict(b=10),
                        font=dict(size=13, color="#aaaaaa"),
                    ),
                    yaxis=dict(
                        title="% Change",
                        showgrid=True,
                        gridcolor=_grid_color,
                        ticksuffix="%",
                    ),
                    xaxis=dict(showgrid=True, gridcolor=_grid_color),
                    plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)",
                    legend=dict(
                        orientation="h",
                        yanchor="bottom",
                        y=1.02,
                        xanchor="right",
                        x=0.99,
                        font=dict(size=11),
                        bgcolor="rgba(0,0,0,0)",
                    ),
                    hovermode="x unified",
                )
                st.plotly_chart(fig_comp, width="stretch")

                for row_start in range(0, len(symbols), 2):
                    cols = st.columns(min(2, len(symbols) - row_start))
                    for col_idx, sym in enumerate(symbols[row_start : row_start + 2]):
                        df = dfs.get(sym)
                        with cols[col_idx]:
                            if df is not None and not df.empty:
                                fig_mini = go.Figure()
                                fig_mini.add_trace(
                                    go.Candlestick(
                                        x=df["datetime"],
                                        open=df["open"],
                                        high=df["high"],
                                        low=df["low"],
                                        close=df["close"],
                                        name=sym,
                                        showlegend=False,
                                        increasing_line_color="#00e676",
                                        decreasing_line_color="#ff1744",
                                        increasing_fillcolor="#00e676",
                                        decreasing_fillcolor="#ff1744",
                                    )
                                )
                                fig_mini.update_layout(
                                    template=_plotly_template,
                                    height=280,
                                    margin=dict(l=0, r=0, t=30, b=0),
                                    title=dict(text=sym, font=dict(size=13)),
                                    xaxis=dict(showgrid=False, showticklabels=False),
                                    yaxis=dict(
                                        showgrid=True,
                                        gridcolor=_grid_color,
                                        tickfont=dict(size=10),
                                    ),
                                    xaxis_rangeslider_visible=False,
                                    plot_bgcolor="#0e1117",
                                    paper_bgcolor="#0e1117",
                                )
                                st.plotly_chart(fig_mini, width="stretch")
                            else:
                                st.info(f"⏳ {sym} — waiting for data…")

            # ── Order books ──────────────────────────────────────────────────────
            st.divider()
            st.markdown("### 📖 Order Book")
            for _ob_idx, sym in enumerate(symbols):
                st.markdown(f"#### {sym}")
                ob_stream = st.session_state.get(f"ob_{sym}")
                if not isinstance(ob_stream, BinanceOrderBookStream):
                    st.info(f"⏳ {sym} — initializing order book stream…")
                    continue

                if not ob_stream.book.has_data:
                    st.info(f"⏳ {sym} — waiting for order book data…")
                    continue

                book = ob_stream.book
                bids, asks = book.snapshot()
                m1, m2, m3, m4 = st.columns(4)
                m1.metric(
                    "Best Bid", f"${book.best_bid:,.2f}" if book.best_bid else "—"
                )
                m2.metric(
                    "Best Ask", f"${book.best_ask:,.2f}" if book.best_ask else "—"
                )
                m3.metric("Spread", f"${book.spread:,.2f}" if book.spread else "—")
                m4.metric(
                    "Spread %", f"{book.spread_pct:.4f}%" if book.spread_pct else "—"
                )

                top_n = 10
                bid_df = pd.DataFrame(bids[:top_n], columns=["Price", "Qty"])
                ask_df = pd.DataFrame(asks[:top_n], columns=["Price", "Qty"])
                bid_df["Total (USDT)"] = (bid_df["Price"] * bid_df["Qty"]).round(2)
                ask_df["Total (USDT)"] = (ask_df["Price"] * ask_df["Qty"]).round(2)

                def _green_shade(col):
                    mx = col.max() or 1
                    return [
                        f"background-color: rgba(0,230,118,{v / mx * 0.55:.2f})"
                        for v in col
                    ]

                def _red_shade(col):
                    mx = col.max() or 1
                    return [
                        f"background-color: rgba(255,23,68,{v / mx * 0.45:.2f})"
                        for v in col
                    ]

                ob_col1, ob_col2 = st.columns(2)
                with ob_col1:
                    st.markdown("**Bids 🟢** — buyers")
                    st.dataframe(
                        bid_df.style.format(
                            {
                                "Price": "{:.4f}",
                                "Qty": "{:.5f}",
                                "Total (USDT)": "{:,.2f}",
                            }
                        ).apply(_green_shade, subset=["Total (USDT)"]),
                        width="stretch",
                        hide_index=True,
                    )
                with ob_col2:
                    st.markdown("**Asks 🔴** — sellers")
                    st.dataframe(
                        ask_df.style.format(
                            {
                                "Price": "{:.4f}",
                                "Qty": "{:.5f}",
                                "Total (USDT)": "{:,.2f}",
                            }
                        ).apply(_red_shade, subset=["Total (USDT)"]),
                        width="stretch",
                        hide_index=True,
                    )

                with st.expander(f"Market Depth Chart — {sym}", expanded=False):
                    bid_prices = [b[0] for b in bids]
                    ask_prices = [a[0] for a in asks]
                    bid_cum, ask_cum, cum = [], [], 0
                    for b in bids:
                        cum += b[1]
                        bid_cum.append(cum)
                    cum = 0
                    for a in asks:
                        cum += a[1]
                        ask_cum.append(cum)
                    fig_depth = go.Figure()
                    fig_depth.add_trace(
                        go.Scatter(
                            x=bid_prices,
                            y=bid_cum,
                            mode="lines",
                            name="Bids",
                            line=dict(color="#00e676", width=2),
                            fill="tozeroy",
                            fillcolor="rgba(0,230,118,0.1)",
                        )
                    )
                    fig_depth.add_trace(
                        go.Scatter(
                            x=ask_prices,
                            y=ask_cum,
                            mode="lines",
                            name="Asks",
                            line=dict(color="#ff1744", width=2),
                            fill="tozeroy",
                            fillcolor="rgba(255,23,68,0.1)",
                        )
                    )
                    fig_depth.update_layout(
                        template=_plotly_template,
                        height=300,
                        margin=dict(l=0, r=0, t=20, b=0),
                        xaxis=dict(
                            title="Price (USDT)", showgrid=True, gridcolor=_grid_color
                        ),
                        yaxis=dict(
                            title="Cumulative Volume",
                            showgrid=True,
                            gridcolor=_grid_color,
                        ),
                        plot_bgcolor="rgba(0,0,0,0)",
                        paper_bgcolor="rgba(0,0,0,0)",
                        legend=dict(orientation="h"),
                    )
                    st.plotly_chart(fig_depth, width="stretch")

                if _ob_idx < len(symbols) - 1:
                    st.divider()

        _dashboard()

# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — CHATBOT  (fragment so reruns never touch the active tab)
# ════════════════════════════════════════════════════════════════════════════════

with tab_research:
    if True:  # tab_research
        # Conversation controls — previously in the sidebar, relocated here so they
        # only render when the user is on this tab (st.tabs can't tell us that
        # from the sidebar itself).
        _has_messages = bool(st.session_state.get("chat_messages"))
        _conv_header_col, _new_chat_col, _clear_all_col = st.columns([3, 1, 1])
        with _conv_header_col:
            st.markdown(
                "<div style='font-size:0.72rem; font-weight:700; letter-spacing:0.08em; "
                "color:rgba(255,255,255,0.55); padding-top:6px;'>CONVERSATIONS</div>",
                unsafe_allow_html=True,
            )
        with _new_chat_col:
            if st.button(
                "＋ New Chat",
                use_container_width=True,
                key="new_chat_btn",
                disabled=not _has_messages,
            ):
                _archive_current_chat()
                st.session_state["chat_messages"] = []
                st.session_state["active_conv_id"] = None
                st.rerun()
        with _clear_all_col:
            if st.button(
                "🗑 Clear All",
                use_container_width=True,
                key="clear_all_btn",
            ):
                st.session_state["conversations"] = []
                st.session_state["chat_messages"] = []
                st.session_state["active_conv_id"] = None
                st.rerun()

        _convs = st.session_state["conversations"]
        if not _convs and not _has_messages:
            st.caption("Start chatting — your conversations will appear here.")
        else:
            with st.expander(
                f"Saved conversations ({len(_convs)})",
                expanded=False,
            ):
                _active_id = st.session_state.get("active_conv_id")
                if _has_messages and _active_id is None:
                    st.markdown(
                        "<div style='font-size:0.73rem; color:rgba(255,255,255,0.6); "
                        "padding:6px 10px; border:1px solid rgba(255,255,255,0.12); "
                        "border-radius:6px; margin-bottom:6px;'>▶ Current chat (unsaved)</div>",
                        unsafe_allow_html=True,
                    )
                for _conv in reversed(_convs):
                    _is_active = _conv["id"] == _active_id
                    _label = ("▶ " if _is_active else "") + _conv["title"]
                    if st.button(
                        _label,
                        key=f"conv_{_conv['id']}",
                        use_container_width=True,
                        help=f"{len(_conv['messages'])} messages",
                    ):
                        _archive_current_chat()
                        st.session_state["chat_messages"] = [
                            dict(m) for m in _conv["messages"]
                        ]
                        st.session_state["active_conv_id"] = _conv["id"]
                        st.rerun()

        st.divider()

        @st.fragment
        def _chatbot(sym: str) -> None:
            """Isolated fragment — all reruns stay inside this tab."""
            api_key = (
                os.getenv("OPENAI_API_KEY")
                or os.getenv("gpt_api_key")
                or os.getenv("GEMINI_API_KEY")
            )

            # ── Fixed-height scrollable chat window ───────────────────────────────
            chat_window = st.container(height=560, border=False)
            with chat_window:
                if not st.session_state["chat_messages"]:
                    with st.chat_message("assistant"):
                        st.markdown(
                            f"**AlphaLens Chatbot** — ask anything about "
                            f"**{sym}** or any crypto.\n\n"
                            "I pull live data from **Binance**, **CoinGecko**, **DefiLlama**, "
                            "**Binance Futures**, and **Polymarket** — then give you a cited answer.\n\n"
                            "**Try asking:**\n"
                            f'- *"Is {sym} a good buy right now?"*\n'
                            f'- *"What\'s the technical setup for {sym}?"*\n'
                            f'- *"What do prediction markets say about {sym}?"*\n\n'
                            "Or use the **quick research** presets below."
                        )
                for _mi, msg in enumerate(st.session_state["chat_messages"]):
                    with st.chat_message(msg["role"]):
                        st.markdown(msg["content"].replace("$", r"\$"))
                        if msg["role"] == "assistant" and len(msg["content"]) > 100:
                            _msg_sym = str(msg.get("symbol") or sym)
                            _hist_dl1, _hist_dl2 = st.columns(2)
                            with _hist_dl1:
                                st.download_button(
                                    "⬇️ Download PDF",
                                    data=_report_to_pdf_bytes(msg["content"], _msg_sym),
                                    file_name=f"alphalens-{_msg_sym.lower()}-{_mi}.pdf",
                                    mime="application/pdf",
                                    key=f"dl_hist_pdf_{_mi}",
                                )
                            with _hist_dl2:
                                st.download_button(
                                    "⬇️ Download Word",
                                    data=_report_to_docx_bytes(
                                        msg["content"], _msg_sym
                                    ),
                                    file_name=f"alphalens-{_msg_sym.lower()}-{_mi}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_hist_docx_{_mi}",
                                )

            # Auto-scroll to bottom
            st.components.v1.html(
                """
                <script>
                (function() {
                    var wrappers = window.parent.document.querySelectorAll(
                        '[data-testid="stVerticalBlockBorderWrapper"]'
                    );
                    wrappers.forEach(function(el) { el.scrollTop = el.scrollHeight; });
                })();
                </script>
                """,
                height=0,
            )

            # ── Chat input ────────────────────────────────────────────────────────
            # Both chat_input and quick buttons live outside the fragment.
            # They store into _chat_user_input, consumed here on the next rerun.
            user_typed = st.session_state.pop("_chat_user_input", None)

            prompt_to_run = user_typed

            if prompt_to_run:
                _prompt_sym = sym
                _prompt_upper = prompt_to_run.upper()
                for _coin in _CHATBOT_COIN_OPTIONS:
                    if re.search(rf"\b{re.escape(_coin)}\b", _prompt_upper):
                        _prompt_sym = _coin
                        break
                if not api_key:
                    st.error(
                        "No API key. Set `GEMINI_API_KEY` (free) or `gpt_api_key` in `.env`."
                    )
                else:
                    st.session_state["chat_messages"].append(
                        {
                            "role": "user",
                            "content": prompt_to_run,
                            "symbol": _prompt_sym,
                        }
                    )
                    with st.chat_message("user"):
                        st.markdown(prompt_to_run)

                    with st.chat_message("assistant"):
                        with st.status("🔍 Fetching data…", expanded=True) as status:

                            def _on_tool(name: str, args: dict) -> None:
                                _sym = args.get("symbol", sym)
                                _label = TOOL_LABELS.get(name, name)
                                status.write(f"{_label} for **{_sym}**…")

                            history = [
                                {"role": m["role"], "content": m["content"]}
                                for m in st.session_state["chat_messages"][-10:]
                            ]

                            try:
                                response, tool_log = run_agent(
                                    prompt_to_run,
                                    history=history[:-1],
                                    on_tool_call=_on_tool,
                                )
                            except Exception as e:
                                response = f"Error: {e}"
                                tool_log = []

                            if tool_log:
                                sources = sorted(
                                    set(
                                        t["result"]["source"]
                                        for t in tool_log
                                        if isinstance(t.get("result"), dict)
                                        and "source" in t["result"]
                                    )
                                )
                                status.update(
                                    label=f"✅ Data from: {', '.join(sources)}",
                                    state="complete",
                                    expanded=False,
                                )
                            else:
                                status.update(
                                    label="✅ Done", state="complete", expanded=False
                                )

                        st.markdown(response.replace("$", r"\$"))

                        if response and len(response) > 100:
                            _sym_safe = _prompt_sym.lower()
                            _ts = pd.Timestamp.utcnow().strftime("%Y%m%d-%H%M%S")
                            _dl1, _dl2 = st.columns(2)
                            with _dl1:
                                st.download_button(
                                    "⬇️ Download PDF",
                                    data=_report_to_pdf_bytes(response, _prompt_sym),
                                    file_name=f"alphalens-{_sym_safe}-{_ts}.pdf",
                                    mime="application/pdf",
                                    key=f"dl_pdf_{_ts}",
                                    use_container_width=True,
                                )
                            with _dl2:
                                st.download_button(
                                    "⬇️ Download Word",
                                    data=_report_to_docx_bytes(response, _prompt_sym),
                                    file_name=f"alphalens-{_sym_safe}-{_ts}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_docx_{_ts}",
                                    use_container_width=True,
                                )

                    st.session_state["chat_messages"].append(
                        {
                            "role": "assistant",
                            "content": response,
                            "symbol": _prompt_sym,
                        }
                    )
                    # Cap chat history to avoid unbounded memory growth
                    if len(st.session_state["chat_messages"]) > 200:
                        st.session_state["chat_messages"] = st.session_state[
                            "chat_messages"
                        ][-200:]
                    _archive_current_chat()

        _chatbot(primary.replace("USDT", ""))

        # Quick preset + chat_input must live OUTSIDE the fragment (Streamlit 1.55+)
        from src.prompts.quick_prompts import QUICK_PROMPTS as _QP

        _sym = primary.replace("USDT", "")
        if st.session_state.get("quick_research_coin") not in _CHATBOT_COIN_OPTIONS:
            st.session_state["quick_research_coin"] = _sym
        if st.session_state.pop("_reset_quick_research_dropdown", False):
            st.session_state["quick_research_dropdown"] = _QUICK_RESEARCH_PLACEHOLDER
        _qr_options = [_QUICK_RESEARCH_PLACEHOLDER] + [
            f"{_e} {_l}" for _l, _e, _ in _QP
        ]
        _qr_col, _coin_col = st.columns(2)
        with _qr_col:
            st.selectbox(
                "Quick research",
                _qr_options,
                key="quick_research_dropdown",
                help="Pick a preset to fill the chat with a structured question.",
            )
        with _coin_col:
            st.selectbox(
                "Cryptocurrency",
                _CHATBOT_COIN_OPTIONS,
                key="quick_research_coin",
                help="Pick which cryptocurrency quick research should be about.",
            )
        _quick_ready = (
            st.session_state.get("quick_research_dropdown")
            != _QUICK_RESEARCH_PLACEHOLDER
        )
        st.button(
            "Generate quick research",
            use_container_width=True,
            on_click=_run_selected_quick_research,
            disabled=not _quick_ready,
            type="primary" if _quick_ready else "secondary",
        )

        _typed = st.chat_input(f"Ask about {_sym} or any crypto…")
        if _typed:
            st.session_state["_chat_user_input"] = _typed
            st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# TAB — SIGNAL SCANNER
# ════════════════════════════════════════════════════════════════════════════════

with tab_signal_scanner:
    if True:  # tab_signal_scanner
        # Scanner controls — relocated from the sidebar so they render only when
        # this tab is active.
        _ss_lbl = _active_theme["label"]
        st.markdown(
            f"<div style='font-size:0.72rem; font-weight:700; letter-spacing:0.08em; "
            f"color:{_ss_lbl}; margin-bottom:4px;'>SIGNAL SCANNER CONTROLS</div>",
            unsafe_allow_html=True,
        )
        _sc_tf_col, _sc_btn_col, _sc_auto_col = st.columns([2, 1, 1])
        with _sc_tf_col:
            _siv = st.session_state.get("scanner_interval", "5m")
            _siv_ix = (
                _DASH_INTERVAL_OPTS.index(_siv) if _siv in _DASH_INTERVAL_OPTS else 2
            )
            st.selectbox(
                "Scanner timeframe",
                _DASH_INTERVAL_OPTS,
                index=_siv_ix,
                key="scanner_interval",
                help="Kline size for RSI, MACD, Bollinger, and MA trend in the table.",
            )
        with _sc_btn_col:
            st.markdown("<div style='height:28px;'></div>", unsafe_allow_html=True)
            if st.button(
                "🔄 Refresh",
                use_container_width=True,
                key="scanner_refresh_btn",
                help="Pull latest candles and recompute every row now (ignores the 30s cache).",
            ):
                st.session_state["_scanner_force_fetch"] = True
                st.rerun()
        with _sc_auto_col:
            st.markdown("<div style='height:34px;'></div>", unsafe_allow_html=True)
            st.toggle(
                "Auto-refresh (30 s)",
                key="scanner_auto_refresh",
                help="On this tab, reloads the scanner on a timer.",
            )
        st.caption(
            "All listed assets use one timeframe. Data comes from the local store; "
            "missing history backfills once per symbol."
        )
        st.divider()
        _scan_auto = bool(st.session_state.get("scanner_auto_refresh", True))

        @st.fragment(run_every=30 if _scan_auto else None)
        def _signal_scanner_body():
            from datetime import datetime, timezone

            st.markdown("### 🔍 Signal Scanner")
            iv = str(st.session_state.get("scanner_interval", "5m"))
            force = bool(st.session_state.pop("_scanner_force_fetch", False))

            cached = _scanner_get_fresh_cache(iv, force)
            if cached:
                cache = cached
            else:
                _load_msg = st.empty()
                _load_msg.info("⚡ Loading signals for all assets…")
                progress = st.progress(0)
                loaded_snap = frozenset(st.session_state.get("loaded_history", set()))
                n_sym = len(AVAILABLE_SYMBOLS)
                results: dict = {}
                new_backfill_keys: set = set()
                with ThreadPoolExecutor(max_workers=4) as executor:
                    futures = {
                        executor.submit(
                            _scanner_fetch_symbol_worker,
                            sym,
                            iv,
                            SCANNER_DATA_LIMIT,
                            loaded_snap,
                        ): sym
                        for sym in AVAILABLE_SYMBOLS
                    }
                    completed = 0
                    for future in as_completed(futures):
                        sym_r, df_r, bf_key = future.result()
                        results[sym_r] = df_r
                        if bf_key:
                            new_backfill_keys.add(bf_key)
                        completed += 1
                        progress.progress(min(1.0, completed / n_sym) if n_sym else 1.0)

                if new_backfill_keys:
                    st.session_state.setdefault("loaded_history", set()).update(
                        new_backfill_keys
                    )

                _load_msg.empty()
                progress.empty()

                rows_out: list[dict] = []
                for sym in AVAILABLE_SYMBOLS:
                    df = results.get(sym)
                    if df is None:
                        rows_out.append({"symbol": sym, "loading": True})
                        continue
                    last = df.iloc[-1]
                    prev = df.iloc[-2]
                    close_f = float(last["close"])
                    prev_c = float(prev["close"])
                    chg_pct = ((close_f - prev_c) / prev_c) * 100 if prev_c else 0.0

                    rsi = last.get("rsi")
                    macd = last.get("macd")
                    sig = last.get("macd_signal")
                    sma20 = last.get("sma_20")
                    sma50 = last.get("sma_50")
                    bb_u = last.get("bb_upper")
                    bb_l = last.get("bb_lower")

                    bb_pct = None
                    if (
                        bb_u is not None
                        and bb_l is not None
                        and not pd.isna(bb_u)
                        and not pd.isna(bb_l)
                    ):
                        bw = float(bb_u) - float(bb_l)
                        if bw > 0:
                            bb_pct = (close_f - float(bb_l)) / bw * 100.0

                    bull_c, bear_c, _neut_c = _scanner_signal_counts(
                        rsi, macd, sig, bb_pct, sma20, sma50
                    )
                    overall = _scanner_overall_label(bull_c, bear_c)
                    summ = indicator_summary(df)

                    rows_out.append(
                        {
                            "symbol": sym,
                            "loading": False,
                            "close": close_f,
                            "chg_pct": chg_pct,
                            "rsi": rsi,
                            "macd": macd,
                            "sig": sig,
                            "bb_pct": bb_pct,
                            "sma20": sma20,
                            "sma50": sma50,
                            "overall": overall,
                            "macd_hist": last.get("macd_hist"),
                            "macd_note": summ.get("macd_crossover", ""),
                        }
                    )

                ts_iso = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
                _scanner_set_display_cache(iv, rows_out, ts_iso)
                cache = st.session_state["scanner_cache"]

            rows = list(cache.get("rows") or [])
            ts_str = cache.get("ts_str", "—")
            _hdr_muted = "#5b6470" if _is_light_theme else "#b8c4d0"
            st.markdown(
                f'<p style="font-size:13px;line-height:1.45;color:{_hdr_muted};margin:0 0 6px 0;">'
                f"Live indicator snapshot across all tracked assets · {html.escape(iv)} candles · "
                f"Updated {html.escape(ts_str)}</p>",
                unsafe_allow_html=True,
            )
            _leg_muted = _hdr_muted
            st.markdown(
                f'<p style="font-size:13px;line-height:1.45;color:{_leg_muted};margin:0 0 16px 0;">'
                "<strong>Legend:</strong> Green = bullish signal · Red = bearish signal · "
                "Gray = neutral</p>",
                unsafe_allow_html=True,
            )

            t = _active_theme
            g_txt = t["green"]
            r_txt = t["red"]
            g_cell = (
                "rgba(22,163,74,0.12)" if _is_light_theme else "rgba(34,197,94,0.15)"
            )
            r_cell = (
                "rgba(220,38,38,0.12)" if _is_light_theme else "rgba(239,68,68,0.15)"
            )
            n_cell = "rgba(0,0,0,0.04)" if _is_light_theme else "rgba(255,255,255,0.05)"
            n_txt = "#6b7280" if _is_light_theme else "#9ca3af"
            desc_clr = "#5b6470" if _is_light_theme else "#b8c4d0"

            def _rsi_sub_color(rv):
                if rv is None or pd.isna(rv):
                    return desc_clr
                v = float(rv)
                if v < 30:
                    return g_txt
                if v > 70:
                    return r_txt
                return desc_clr

            def _macd_sub_color(lb_m: str):
                if lb_m == "bullish":
                    return g_txt
                if lb_m == "bearish":
                    return r_txt
                return desc_clr

            def _bb_sub_color(lb_b: str):
                if lb_b == "near bottom":
                    return g_txt
                if lb_b == "near top":
                    return r_txt
                return desc_clr

            def _rsi_style(rv):
                if rv is None or pd.isna(rv):
                    return n_cell, n_txt, "neutral"
                v = float(rv)
                if v < 30:
                    return g_cell, g_txt, "oversold"
                if v > 70:
                    return r_cell, r_txt, "overbought"
                return n_cell, n_txt, "neutral"

            def _macd_style(mv, sv):
                if mv is None or sv is None or pd.isna(mv) or pd.isna(sv):
                    return n_cell, n_txt, "neutral"
                if float(mv) > float(sv):
                    return g_cell, g_txt, "bullish"
                if float(mv) < float(sv):
                    return r_cell, r_txt, "bearish"
                return n_cell, n_txt, "neutral"

            def _bb_style(pv):
                if pv is None or pd.isna(pv):
                    return n_cell, n_txt, "mid-band"
                v = float(pv)
                if v < 20:
                    return g_cell, g_txt, "near bottom"
                if v > 80:
                    return r_cell, r_txt, "near top"
                return n_cell, n_txt, "mid-band"

            def _trend_html(s20, s50):
                if (
                    s20 is None
                    or s50 is None
                    or pd.isna(s20)
                    or pd.isna(s50)
                    or float(s50) == 0
                ):
                    return f'<span class="alphalens-scanner-trend" style="color:{n_txt};">→ Neutral</span>'
                a20, a50 = float(s20), float(s50)
                if abs(a20 - a50) / abs(a50) < 0.001:
                    return f'<span class="alphalens-scanner-trend" style="color:{n_txt};">→ Neutral</span>'
                if a20 > a50:
                    return f'<span class="alphalens-scanner-trend" style="color:{g_txt};">↑ Bullish</span>'
                return f'<span class="alphalens-scanner-trend" style="color:{r_txt};">↓ Bearish</span>'

            th = (
                "<thead><tr>"
                "<th>Asset</th><th>Price</th><th>RSI (14)</th><th>MACD</th>"
                "<th>BB Position</th><th>Trend</th><th>Overall Signal</th>"
                "</tr></thead>"
            )
            tb_parts: list[str] = []
            for rw in rows:
                sym = rw["symbol"]
                tick = sym.replace("USDT", "")
                badge = SCANNER_COIN_BADGE_BG.get(tick, SCANNER_BADGE_DEFAULT)
                if rw.get("loading"):
                    tb_parts.append(
                        "<tr class='alphalens-scanner-row'>"
                        f'<td><span style="display:inline-flex;align-items:center;justify-content:center;'
                        f"min-width:40px;padding:4px 10px;border-radius:9999px;background:{badge};"
                        f'color:#fff;font-weight:700;font-size:12px;">{html.escape(tick)}</span></td>'
                        f'<td colspan="6" style="color:{desc_clr};font-size:13px;font-style:italic;">'
                        "loading…</td></tr>"
                    )
                    continue

                chg = rw["chg_pct"]
                chg_c = g_txt if chg >= 0 else r_txt
                rsi = rw["rsi"]
                bg_rsi, fg_rsi, lb_rsi = _rsi_style(rsi)
                macd = rw["macd"]
                sigv = rw["sig"]
                bg_m, fg_m, lb_m = _macd_style(macd, sigv)
                bbp = rw["bb_pct"]
                bg_b, fg_b, lb_b = _bb_style(bbp)

                rsi_disp = (
                    f"{float(rsi):.1f}" if rsi is not None and not pd.isna(rsi) else "—"
                )
                macd_disp = (
                    f"{float(macd):.4f}"
                    if macd is not None and not pd.isna(macd)
                    else "—"
                )
                bb_disp = (
                    f"{float(bbp):.0f}%"
                    if bbp is not None and not pd.isna(bbp)
                    else "—"
                )

                sub_rsi = _rsi_sub_color(rsi)
                sub_m = _macd_sub_color(lb_m)
                sub_b = _bb_sub_color(lb_b)
                tb_parts.append(
                    "<tr class='alphalens-scanner-row'>"
                    f'<td><span style="display:inline-flex;align-items:center;justify-content:center;'
                    f"min-width:40px;padding:4px 10px;border-radius:9999px;background:{badge};"
                    f'color:#fff;font-weight:700;font-size:12px;">{html.escape(tick)}</span></td>'
                    f"<td><div class='alphalens-scanner-cell-main' style='color:{t['text']};'>"
                    f"${rw['close']:,.2f}</div>"
                    f"<div class='alphalens-scanner-cell-sub' style='color:{chg_c};'>"
                    f"{chg:+.2f}%</div></td>"
                    f"<td style='background:{bg_rsi};color:{fg_rsi};'>"
                    f"<div class='alphalens-scanner-cell-main'>{rsi_disp}</div>"
                    f"<div class='alphalens-scanner-cell-sub' style='color:{sub_rsi};'>{lb_rsi}</div></td>"
                    f"<td style='background:{bg_m};color:{fg_m};'>"
                    f"<div class='alphalens-scanner-cell-main'>{macd_disp}</div>"
                    f"<div class='alphalens-scanner-cell-sub' style='color:{sub_m};'>{lb_m}</div></td>"
                    f"<td style='background:{bg_b};color:{fg_b};'>"
                    f"<div class='alphalens-scanner-cell-main'>{bb_disp}</div>"
                    f"<div class='alphalens-scanner-cell-sub' style='color:{sub_b};'>{lb_b}</div></td>"
                    f"<td>{_trend_html(rw['sma20'], rw['sma50'])}</td>"
                    f"<td class='alphalens-scanner-cell-main' style='font-weight:700;color:{t['text']};'>"
                    f"{rw['overall']}</td>"
                    "</tr>"
                )

            st.markdown(
                '<div class="alphalens-scanner-wrap">'
                '<table class="alphalens-scanner-table">'
                + th
                + "<tbody>"
                + "".join(tb_parts)
                + "</tbody></table></div>",
                unsafe_allow_html=True,
            )

            done_rows = [r for r in rows if not r.get("loading")]
            n_bull = sum(
                1 for r in done_rows if r.get("overall") in ("🟢 Strong Buy", "🟡 Buy")
            )
            n_bear = sum(
                1
                for r in done_rows
                if r.get("overall") in ("🔴 Strong Sell", "🟡 Sell")
            )
            n_neu = sum(1 for r in done_rows if r.get("overall") == "⚪ Neutral")
            n_str = sum(
                1
                for r in done_rows
                if r.get("overall") in ("🟢 Strong Buy", "🔴 Strong Sell")
            )
            st.markdown(
                f'<div class="alphalens-scanner-summary">'
                f'<span class="alphalens-scanner-pill" style="color:{g_txt};">🟢 {n_bull} Bullish</span>'
                f'<span class="alphalens-scanner-pill" style="color:{r_txt};">🔴 {n_bear} Bearish</span>'
                f'<span class="alphalens-scanner-pill" style="color:{n_txt};">⚪ {n_neu} Neutral</span>'
                f'<span class="alphalens-scanner-pill" style="color:{t["accent_text"]};">'
                f"⚡ {n_str} Strong Signals</span></div>",
                unsafe_allow_html=True,
            )

            oversold = [
                r
                for r in done_rows
                if r.get("rsi") is not None
                and not pd.isna(r["rsi"])
                and float(r["rsi"]) < 35
            ]
            overbought = [
                r
                for r in done_rows
                if r.get("rsi") is not None
                and not pd.isna(r["rsi"])
                and float(r["rsi"]) > 65
            ]
            macd_div = [
                r
                for r in done_rows
                if r.get("macd_hist") is not None and not pd.isna(r["macd_hist"])
            ]

            opp_blocks: list[str] = []
            if oversold:
                best = min(oversold, key=lambda x: float(x["rsi"]))
                tk = best["symbol"].replace("USDT", "")
                bg = SCANNER_COIN_BADGE_BG.get(tk, SCANNER_BADGE_DEFAULT)
                rv = float(best["rsi"])
                opp_blocks.append(
                    f'<div style="background:{t["card"]};border:1px solid {t["border"]};'
                    f'border-radius:{t["radius"]};padding:14px 18px;margin-bottom:10px;">'
                    f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:8px;">'
                    f'<span style="display:inline-flex;min-width:38px;padding:3px 9px;'
                    f"border-radius:9999px;background:{bg};color:#fff;font-weight:700;"
                    f'font-size:11px;">{html.escape(tk)}</span>'
                    f'<span style="font-weight:700;color:{g_txt};">Most oversold</span></div>'
                    f'<div style="font-size:1.25rem;font-weight:600;color:{t["text"]};">'
                    f"RSI {rv:.1f}</div>"
                    f'<p style="margin:8px 0 0;font-size:12px;color:{desc_clr};line-height:1.4;">'
                    f"RSI at {rv:.1f} — historically oversold territory.</p></div>"
                )
            if overbought:
                worst = max(overbought, key=lambda x: float(x["rsi"]))
                tk = worst["symbol"].replace("USDT", "")
                bg = SCANNER_COIN_BADGE_BG.get(tk, SCANNER_BADGE_DEFAULT)
                rv = float(worst["rsi"])
                opp_blocks.append(
                    f'<div style="background:{t["card"]};border:1px solid {t["border"]};'
                    f'border-radius:{t["radius"]};padding:14px 18px;margin-bottom:10px;">'
                    f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:8px;">'
                    f'<span style="display:inline-flex;min-width:38px;padding:3px 9px;'
                    f"border-radius:9999px;background:{bg};color:#fff;font-weight:700;"
                    f'font-size:11px;">{html.escape(tk)}</span>'
                    f'<span style="font-weight:700;color:{r_txt};">Most overbought</span></div>'
                    f'<div style="font-size:1.25rem;font-weight:600;color:{t["text"]};">'
                    f"RSI {rv:.1f}</div>"
                    f'<p style="margin:8px 0 0;font-size:12px;color:{desc_clr};line-height:1.4;">'
                    f"RSI at {rv:.1f} — stretched toward overbought.</p></div>"
                )
            if macd_div:
                best_m = max(macd_div, key=lambda x: abs(float(x["macd_hist"])))
                tk = best_m["symbol"].replace("USDT", "")
                bg = SCANNER_COIN_BADGE_BG.get(tk, SCANNER_BADGE_DEFAULT)
                hv = float(best_m["macd_hist"])
                _mn = html.escape(str(best_m.get("macd_note") or ""))
                opp_blocks.append(
                    f'<div style="background:{t["card"]};border:1px solid {t["border"]};'
                    f'border-radius:{t["radius"]};padding:14px 18px;margin-bottom:10px;">'
                    f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:8px;">'
                    f'<span style="display:inline-flex;min-width:38px;padding:3px 9px;'
                    f"border-radius:9999px;background:{bg};color:#fff;font-weight:700;"
                    f'font-size:11px;">{html.escape(tk)}</span>'
                    f'<span style="font-weight:700;color:{t["accent_text"]};">'
                    f"Strongest MACD histogram</span></div>"
                    f'<div style="font-size:1.25rem;font-weight:600;color:{t["text"]};">'
                    f"Hist {hv:+.4f}</div>"
                    f'<p style="margin:8px 0 0;font-size:12px;color:{desc_clr};line-height:1.4;">'
                    f"{_mn or 'Largest histogram reading in this scan — momentum stands out vs peers.'}"
                    "</p></div>"
                )

            if opp_blocks:
                st.markdown(
                    '<div style="font-size:0.72rem;font-weight:700;letter-spacing:0.08em;'
                    f'color:{t["label"]};margin:20px 0 10px 0;">TOP OPPORTUNITIES</div>',
                    unsafe_allow_html=True,
                )
                st.markdown("".join(opp_blocks), unsafe_allow_html=True)

        _signal_scanner_body()

# ════════════════════════════════════════════════════════════════════════════════
# TAB — LIVE PREDICTION MARKETS
# ════════════════════════════════════════════════════════════════════════════════

with tab_prediction:
    if True:  # tab_prediction
        st.markdown("### 🎯 Live Prediction Markets")
        st.caption("Auto-refreshing every 30s · Polymarket")

        # Market filters — relocated from the sidebar.
        _pm_coin_col, _pm_time_col = st.columns(2)
        with _pm_coin_col:
            st.multiselect(
                "Filter by coin",
                SCANNER_COIN_OPTIONS,
                default=["All"],
                key="scanner_coin_filter",
            )
        with _pm_time_col:
            _tf = st.session_state.get("scanner_time_filter", "All")
            _tf_ix = (
                SCANNER_TIME_OPTIONS.index(_tf) if _tf in SCANNER_TIME_OPTIONS else 0
            )
            st.selectbox(
                "Time period",
                SCANNER_TIME_OPTIONS,
                index=_tf_ix,
                key="scanner_time_filter",
            )
        st.divider()

        if "mkt_history" not in st.session_state:
            st.session_state["mkt_history"] = {}

        def _fmt_time(hours: float) -> str:
            if hours <= 0:
                return "Expired"
            if hours < 1:
                return f"{int(hours * 60)}m left"
            if hours < 24:
                return f"{int(hours)}h {int((hours % 1) * 60)}m left"
            d = int(hours / 24)
            h = int(hours % 24)
            return f"{d}d {h}h left"

        def _fmt_vol(v: float) -> str:
            if v >= 1_000_000:
                return f"${v / 1_000_000:.1f}M"
            if v >= 1_000:
                return f"${v / 1_000:.0f}K"
            return f"${v:.0f}"

        @st.fragment(run_every=30 if auto_refresh else None)
        def _live_markets():
            from datetime import datetime, timezone
            from collections import defaultdict
            from src.polymarket import fetch_crypto_markets

            coin_filter = st.session_state.get("scanner_coin_filter", ["All"])
            time_filter = st.session_state.get("scanner_time_filter", "All")

            def _hours_until(expiry):
                if expiry is None:
                    return 168.0
                delta = (expiry - datetime.now(timezone.utc)).total_seconds() / 3600
                return max(0, delta)

            try:
                raw_markets = fetch_crypto_markets()
            except Exception:
                raw_markets = []

            # Only keep markets for coins in our tracked asset list
            tracked_symbols = set(AVAILABLE_SYMBOLS)
            markets = []
            for m in raw_markets:
                if m["symbol"] not in tracked_symbols:
                    continue
                vol = m.get("volume", 0)
                if vol < 10_000:
                    continue
                odds = m["yes_price"] * 100
                # Skip effectively resolved markets (odds pinned near 0% or 100%)
                if odds <= 1 or odds >= 99:
                    continue
                hours = _hours_until(m.get("expiry"))
                markets.append(
                    {
                        "id": m.get("id", ""),
                        "question": m["question"],
                        "symbol": m["symbol"],
                        "market_odds": round(odds, 1),
                        "volume": vol,
                        "liquidity": m.get("liquidity", 0),
                        "hours_left": round(hours, 1),
                        "expiry": m.get("expiry"),
                        "slug": m.get("slug", ""),
                        "clob_token_id": m.get("clob_token_id"),
                        "threshold": m.get("threshold"),
                        "direction": m.get("direction", "above"),
                        "event_id": m.get("event_id", ""),
                        "event_title": m.get("event_title", ""),
                    }
                )

            if not markets:
                st.info(
                    "No crypto prediction markets with >$10K volume found. "
                    "Markets may be inactive or APIs temporarily unreachable."
                )
                return

            # Track current odds for all markets (cheap — no HTTP calls)
            now = datetime.now(timezone.utc)
            for mkt in markets:
                key = mkt["id"] or mkt["question"][:60]
                if key not in st.session_state["mkt_history"]:
                    st.session_state["mkt_history"][key] = []
                st.session_state["mkt_history"][key].append(
                    (now.timestamp(), mkt["market_odds"])
                )
                st.session_state["mkt_history"][key] = st.session_state["mkt_history"][
                    key
                ][-60:]

            # Prune stale market keys no longer in the active set
            active_keys = {m["id"] or m["question"][:60] for m in markets}
            stale = [k for k in st.session_state["mkt_history"] if k not in active_keys]
            for k in stale:
                del st.session_state["mkt_history"][k]

            # ── Filter by coin ────────────────────────────────────────────────
            selected_coins = coin_filter if coin_filter else ["All"]
            if "All" not in selected_coins:
                filter_pairs = {c + "USDT" for c in selected_coins}
                visible = [m for m in markets if m["symbol"] in filter_pairs]
            else:
                visible = list(markets)

            # ── Filter by time period ─────────────────────────────────────────
            if time_filter != "All":
                time_ranges = {
                    "Today": 24,
                    "This Week": 168,
                    "This Month": 744,
                    "Long-term": float("inf"),
                }
                max_h = time_ranges[time_filter]
                if time_filter == "Long-term":
                    visible = [m for m in visible if m["hours_left"] > 744]
                else:
                    prev_max = {"Today": 0, "This Week": 24, "This Month": 168}[
                        time_filter
                    ]
                    visible = [
                        m for m in visible if prev_max < m["hours_left"] <= max_h
                    ]

            # ── Sort by volume (highest first) ────────────────────────────────
            visible.sort(key=lambda o: o.get("volume", 0), reverse=True)

            # ── Summary stats ─────────────────────────────────────────────────
            total_vol = sum(o.get("volume", 0) for o in visible)
            _selected_specific = [c for c in selected_coins if c != "All"]
            _visible_coin_set = {
                o["symbol"].replace("USDT", "")
                for o in visible
                if o["symbol"] != "CRYPTOUSDT"
            }
            if _selected_specific:
                coins_in_view = _selected_specific
            else:
                coins_in_view = sorted(_visible_coin_set)
            s1, s2, s3 = st.columns(3)
            s1.metric("Markets", len(visible))
            s2.metric("Total Volume", _fmt_vol(total_vol))
            s3.metric("Coins", " · ".join(coins_in_view) if coins_in_view else "—")

            _missing_selected = [
                c for c in _selected_specific if c not in _visible_coin_set
            ]
            if _missing_selected:
                st.info(
                    f"No prediction market data for {', '.join(_missing_selected)}. "
                    "Try a different coin."
                )
            if not visible:
                return

            # ── Implied price range summary ────────────────────────────────────
            coin_targets: dict[str, list[dict]] = defaultdict(list)
            for mkt in visible:
                if mkt.get("threshold") and mkt["symbol"] != "CRYPTOUSDT":
                    coin_targets[mkt["symbol"]].append(mkt)

            if coin_targets:
                from datetime import datetime, timedelta, timezone as _tz

                _now = datetime.now(_tz.utc)

                consensus_rows: list[dict] = []
                for sym_key in sorted(
                    coin_targets,
                    key=lambda s: sum(m["volume"] for m in coin_targets[s]),
                    reverse=True,
                )[:4]:
                    coin_name = sym_key.replace("USDT", "")
                    targets = coin_targets[sym_key]

                    # Group by event (= same resolution date), pick highest-volume event
                    by_event: dict[str, list[dict]] = defaultdict(list)
                    for t in targets:
                        eid = t.get("event_id") or "unknown"
                        by_event[eid].append(t)
                    top_event = max(
                        by_event.values(), key=lambda g: sum(m["volume"] for m in g)
                    )

                    # Pick the market with the highest probability
                    best = max(top_event, key=lambda t: t["market_odds"])

                    # Resolution date
                    h = best["hours_left"]
                    exp_dt = _now + timedelta(hours=h)
                    by_when = exp_dt.strftime("by %b %d, %Y")

                    consensus_rows.append(
                        {
                            "ticker": coin_name,
                            "pct": float(best["market_odds"]),
                            "direction": best.get("direction", "above"),
                            "threshold": best["threshold"],
                            "by_when": by_when,
                        }
                    )

                if consensus_rows:
                    _t = _active_theme
                    _desc_clr = "#4b5563" if _is_light_theme else "#d1d5db"
                    _track_clr = "#d1d5db" if _is_light_theme else "#374151"
                    _html_parts: list[str] = [
                        f'<div style="background:{_t["card"]};border:1px solid {_t["border"]};'
                        f'border-top:2px solid #f59e0b;border-radius:{_t["radius"]};margin-bottom:16px;">'
                        '<div style="display:flex;align-items:center;justify-content:space-between;'
                        'flex-wrap:wrap;gap:8px;padding:10px 16px 6px 16px;">'
                        '<div style="display:flex;align-items:center;gap:8px;">'
                        '<span class="alphalens-consensus-live-dot"></span>'
                        f'<span style="font-size:0.72rem;font-weight:600;color:{_t["label"]};'
                        'letter-spacing:0.08em;text-transform:uppercase;">MARKET CONSENSUS</span>'
                        "</div>"
                        '<span style="font-size:10px;color:#6b7280;">Powered by Polymarket</span>'
                        "</div>"
                    ]
                    for crow in consensus_rows:
                        _tick = crow["ticker"]
                        _pct = crow["pct"]
                        _dir = crow["direction"]
                        _thr = crow["threshold"]
                        _by = crow["by_when"]
                        _badge = _CONSENSUS_BADGE_BG.get(_tick, "#64748b")
                        _bar_c = _consensus_prob_bar_color(_pct)
                        _fill = min(100.0, max(0.0, _pct))
                        _line = f"{_pct:.0f}% chance {_dir} ${_thr:,.0f} {_by}"
                        _html_parts.append(
                            '<div class="alphalens-consensus-row">'
                            f'<span style="display:inline-flex;align-items:center;justify-content:center;'
                            f"min-width:38px;padding:3px 9px;border-radius:9999px;background:{_badge};"
                            f'color:#fff;font-weight:700;font-size:11px;line-height:1;flex-shrink:0;">'
                            f"{html.escape(_tick)}</span>"
                            f'<span style="flex:1;min-width:0;font-size:13px;color:{_desc_clr};line-height:1.45;">'
                            f"{html.escape(_line)}</span>"
                            '<div style="display:flex;align-items:center;gap:10px;flex-shrink:0;'
                            'width:128px;max-width:36%;">'
                            f'<div style="flex:1;height:4px;background:{_track_clr};border-radius:2px;'
                            f'overflow:hidden;min-width:52px;">'
                            f'<div style="width:{_fill:.2f}%;height:100%;background:{_bar_c};'
                            f'border-radius:2px;"></div></div>'
                            f'<span style="font-weight:700;font-size:12px;color:{_bar_c};min-width:38px;'
                            f'text-align:right;white-space:nowrap;">'
                            f"{_pct:.0f}%</span></div></div>"
                        )
                    _html_parts.append("</div>")
                    st.markdown("".join(_html_parts), unsafe_allow_html=True)

            # ── Event grouping — top 2 by volume get histogram ─────────────────
            event_groups: dict[str, list[dict]] = defaultdict(list)
            for mkt in visible:
                eid = mkt.get("event_id")
                if eid and mkt.get("threshold"):
                    event_groups[eid].append(mkt)

            sorted_events = sorted(
                event_groups.items(),
                key=lambda x: sum(m["volume"] for m in x[1]),
                reverse=True,
            )
            histogram_ids: set[str] = set()
            histogram_count = 0
            for eid, group in sorted_events:
                if len(group) >= 3 and histogram_count < 2:
                    histogram_count += 1
                    histogram_ids.add(eid)
                    _render_event_group(group)

            # ── Sparkline cards (skip markets already in histograms) ──────────
            card_mkts = [
                m for m in visible if m.get("event_id", "") not in histogram_ids
            ]
            for row_start in range(0, len(card_mkts), 2):
                cols = st.columns(2)
                for col_idx, mkt in enumerate(card_mkts[row_start : row_start + 2]):
                    with cols[col_idx]:
                        _render_market_card(mkt)

        def _render_event_group(group: list[dict]) -> None:
            """Render a price-target event as a distribution chart."""
            title = group[0].get("event_title") or group[0]["question"][:50]
            coin = group[0]["symbol"].replace("USDT", "")
            total_vol = sum(m["volume"] for m in group)
            direction = group[0].get("direction", "above")
            _grp_expiry = group[0].get("expiry")
            _grp_end = (
                f" · Ends {_grp_expiry.strftime('%b %d, %Y')}" if _grp_expiry else ""
            )

            # Sort by threshold
            sorted_g = sorted(group, key=lambda m: m["threshold"] or 0)

            with st.container(border=True):
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:10px;">'
                    f'<span style="background:rgba(124,131,253,0.15);color:#a5abff;'
                    f'padding:2px 8px;border-radius:6px;font-size:0.72rem;font-weight:700;">{coin}</span>'
                    f'<span style="font-weight:600;font-size:0.95rem;">{title}</span>'
                    f'<span style="color:rgba(255,255,255,0.35);font-size:0.75rem;margin-left:auto;">'
                    f"{_fmt_vol(total_vol)} vol · {len(group)} markets{_grp_end} · "
                    f'<span style="color:#2E5CFF;">Polymarket</span></span>'
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # Bar chart — probability at each strike
                labels = []
                probs = []
                colors = []
                for m in sorted_g:
                    t = m["threshold"]
                    if t >= 1000:
                        labels.append(f"${t / 1000:.0f}K")
                    else:
                        labels.append(f"${t:,.0f}")
                    p = m["market_odds"]
                    probs.append(p)
                    colors.append(
                        "#00e676" if p >= 50 else "#ff1744" if p < 20 else "#ff9800"
                    )

                fig = go.Figure()
                fig.add_trace(
                    go.Bar(
                        x=labels,
                        y=probs,
                        marker_color=colors,
                        text=[f"{p:.0f}%" for p in probs],
                        textposition="outside",
                        textfont=dict(color="rgba(255,255,255,0.7)", size=11),
                        hovertemplate="%{x}: %{y:.1f}%<extra></extra>",
                    )
                )
                fig.update_layout(
                    height=200,
                    margin=dict(l=0, r=0, t=10, b=0),
                    xaxis=dict(
                        showgrid=False,
                        color="rgba(255,255,255,0.5)",
                        tickfont=dict(size=11),
                    ),
                    yaxis=dict(
                        visible=False,
                        range=[0, max(probs) * 1.25] if probs else [0, 100],
                    ),
                    plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)",
                    dragmode=False,
                )
                fig.update_xaxes(fixedrange=True)
                fig.update_yaxes(fixedrange=True)
                st.plotly_chart(
                    fig,
                    use_container_width=True,
                    config={"displayModeBar": False},
                    key=f"hist_{group[0]['event_id']}",
                )

        def _render_market_card(opp: dict) -> None:
            odds = opp["market_odds"]
            odds_color = "#00e676" if odds >= 50 else "#ff1744"
            coin = opp["symbol"].replace("USDT", "")

            with st.container(border=True):
                # Header: coin badge + question
                st.markdown(
                    f'<div style="display:flex;align-items:flex-start;gap:10px;">'
                    f'<span style="background:rgba(124,131,253,0.15);color:#a5abff;'
                    f"padding:2px 8px;border-radius:6px;font-size:0.72rem;"
                    f'font-weight:700;letter-spacing:0.5px;white-space:nowrap;margin-top:2px;">{coin}</span>'
                    f'<span style="font-weight:600;font-size:0.92rem;line-height:1.4;">'
                    f"{opp['question']}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # Probability + change (lazy-seed price history on first render)
                key = opp["id"] or opp["question"][:60]
                history = st.session_state.get("mkt_history", {}).get(key, [])
                if len(history) <= 1 and opp.get("clob_token_id"):
                    from src.polymarket import fetch_price_history

                    raw = fetch_price_history(
                        opp["clob_token_id"], interval="1w", fidelity=40
                    )
                    seed = [(t, p * 100) for t, p in raw]
                    if seed:
                        st.session_state["mkt_history"][key] = seed + history
                        history = st.session_state["mkt_history"][key]
                change_html = ""
                if len(history) >= 2:
                    first_p, last_p = history[0][1], history[-1][1]
                    delta = last_p - first_p
                    if abs(delta) >= 0.1:
                        d_color = "#00e676" if delta > 0 else "#ff1744"
                        d_arrow = "+" if delta > 0 else ""
                        change_html = (
                            f'<span style="font-size:0.82rem;font-weight:600;color:{d_color};'
                            f'margin-left:10px;">{d_arrow}{delta:.1f}%</span>'
                        )

                st.markdown(
                    f'<div style="margin:8px 0 2px;">'
                    f'<span style="font-size:1.8rem;font-weight:700;color:{odds_color};">'
                    f"{odds:.0f}%</span>"
                    f'<span style="font-size:0.82rem;color:rgba(255,255,255,0.45);'
                    f'margin-left:6px;font-weight:500;">Chance</span>'
                    f"{change_html}"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # Chart — Polymarket/Robinhood style
                if len(history) >= 2:
                    from datetime import datetime, timezone

                    times = [
                        datetime.fromtimestamp(t, tz=timezone.utc) for t, _ in history
                    ]
                    prices = [p for _, p in history]
                    first_p, last_p = prices[0], prices[-1]
                    line_color = "#00e676" if last_p >= first_p else "#ff1744"
                    fill_top = (
                        "rgba(0,230,118,0.15)"
                        if last_p >= first_p
                        else "rgba(255,23,68,0.15)"
                    )

                    fig = go.Figure()
                    fig.add_trace(
                        go.Scatter(
                            x=times,
                            y=prices,
                            mode="none",
                            fill="tozeroy",
                            fillcolor=fill_top,
                            showlegend=False,
                            hoverinfo="skip",
                        )
                    )
                    fig.add_trace(
                        go.Scatter(
                            x=times,
                            y=prices,
                            mode="lines",
                            line=dict(
                                color=line_color,
                                width=2.5,
                                shape="spline",
                                smoothing=1.0,
                            ),
                            showlegend=False,
                            hovertemplate="%{y:.1f}%<extra></extra>",
                        )
                    )
                    fig.add_trace(
                        go.Scatter(
                            x=[times[-1]],
                            y=[prices[-1]],
                            mode="markers",
                            marker=dict(
                                color=line_color,
                                size=7,
                                line=dict(width=2, color="#0b0e14"),
                            ),
                            showlegend=False,
                            hoverinfo="skip",
                        )
                    )
                    y_min = max(0, min(prices) - 5)
                    y_max = min(100, max(prices) + 5)
                    fig.update_layout(
                        height=110,
                        margin=dict(l=0, r=0, t=0, b=0),
                        xaxis=dict(visible=False, showgrid=False),
                        yaxis=dict(
                            visible=False,
                            showgrid=False,
                            range=[y_min, y_max],
                            fixedrange=True,
                        ),
                        plot_bgcolor="rgba(0,0,0,0)",
                        paper_bgcolor="rgba(0,0,0,0)",
                        hoverlabel=dict(
                            bgcolor="#1a1e2e",
                            bordercolor="rgba(255,255,255,0.1)",
                            font=dict(color="#fff", size=12),
                        ),
                        dragmode=False,
                    )
                    fig.update_xaxes(fixedrange=True)
                    st.plotly_chart(
                        fig,
                        use_container_width=True,
                        config={"displayModeBar": False},
                        key=f"spark_{opp['id']}",
                    )

                # Footer: volume · time left · end date | platform watermark
                _expiry = opp.get("expiry")
                _end_str = _expiry.strftime("%b %d, %Y") if _expiry else ""
                _end_html = f"<span>Ends {_end_str}</span>" if _end_str else ""
                st.markdown(
                    f'<div style="display:flex;justify-content:space-between;'
                    f"align-items:center;margin-top:4px;padding-top:8px;"
                    f'border-top:1px solid rgba(255,255,255,0.06);">'
                    f'<div style="display:flex;gap:14px;font-size:0.75rem;'
                    f'color:rgba(255,255,255,0.4);font-weight:500;">'
                    f"<span>{_fmt_vol(opp['volume'])} vol</span>"
                    f"<span>{_fmt_time(opp['hours_left'])}</span>"
                    f"{_end_html}"
                    f"</div>"
                    f'<span style="font-size:0.68rem;color:rgba(255,255,255,0.25);font-weight:500;'
                    f'letter-spacing:0.03em;color:#2E5CFF;">Polymarket</span>'
                    f"</div>",
                    unsafe_allow_html=True,
                )

        _live_markets()

# ════════════════════════════════════════════════════════════════════════════════
# TAB — NEWS FEED
# ════════════════════════════════════════════════════════════════════════════════

NEWS_COIN_OPTIONS = [
    "All",
    "BTC",
    "ETH",
    "SOL",
    "BNB",
    "XRP",
    "DOGE",
    "ADA",
    "AVAX",
    "LINK",
    "DOT",
    "UNI",
    "NEAR",
    "ARB",
    "SUI",
    "APT",
    "PEPE",
    "TON",
]

with tab_news:
    if True:  # tab_news
        st.session_state["active_tab"] = "News Feed"
        st.markdown("### 📰 Crypto News Feed")
        st.caption(
            "Auto-refreshing every 60s · CoinDesk, CoinTelegraph, The Block, Decrypt, Blockworks, Bitcoin Magazine"
        )

        t = _active_theme
        _news_accent = t["accent"]

        if "news_archive" not in st.session_state:
            st.session_state["news_archive"] = {}

        _nf_col1, _nf_col2, _ = st.columns([1, 1, 2])
        with _nf_col1:
            news_coin_filter = st.selectbox(
                "Filter by coin",
                NEWS_COIN_OPTIONS,
                index=0,
                key="news_coin_filter",
            )
        with _nf_col2:
            news_time_filter = st.selectbox(
                "Time range",
                ["Latest", "Today", "Past 24h", "Past 3 Days", "Past 5 Days"],
                index=0,
                key="news_time_filter",
            )

        @st.fragment(run_every=60 if auto_refresh else None)
        def _news_feed():
            from src.tools import get_crypto_news

            # Always fetch all to build archive, filter for display
            with st.status("Fetching news...", expanded=False) as status:
                result = get_crypto_news("")
                status.update(label="News loaded", state="complete", expanded=False)

            if result.get("error"):
                st.error(f"Failed to fetch news: {result['error']}")
                return

            # Merge into session archive (keyed by url, pruned to 5 days)
            now = time.time()
            five_days_ago = now - (5 * 86400)
            archive = st.session_state["news_archive"]
            for a in result.get("articles", []):
                if a["url"] and a["published_on"] > five_days_ago:
                    archive[a["url"]] = a
            st.session_state["news_archive"] = {
                u: a for u, a in archive.items() if a["published_on"] > five_days_ago
            }

            # Build display list from archive
            all_articles = sorted(
                st.session_state["news_archive"].values(),
                key=lambda a: a["published_on"],
                reverse=True,
            )

            # Time filter
            time_cutoffs = {
                "Latest": 0,
                "Today": now - 86400,
                "Past 24h": now - 86400,
                "Past 3 Days": now - (3 * 86400),
                "Past 5 Days": now - (5 * 86400),
            }
            cutoff = time_cutoffs.get(news_time_filter, 0)
            if cutoff > 0:
                all_articles = [a for a in all_articles if a["published_on"] >= cutoff]

            # Coin filter
            if news_coin_filter != "All":
                kw = news_coin_filter.lower()
                aliases = {
                    "BTC": ["bitcoin", "btc"],
                    "ETH": ["ethereum", "eth", "ether"],
                    "SOL": ["solana", "sol"],
                    "BNB": ["bnb", "binance coin", "binance"],
                    "XRP": ["xrp", "ripple"],
                    "DOGE": ["doge", "dogecoin"],
                    "ADA": ["ada", "cardano"],
                    "AVAX": ["avax", "avalanche"],
                    "DOT": ["dot", "polkadot"],
                    "LINK": ["link", "chainlink"],
                    "UNI": ["uni", "uniswap"],
                    "NEAR": ["near protocol", "near"],
                    "ARB": ["arb", "arbitrum"],
                    "SUI": ["sui"],
                    "APT": ["apt", "aptos"],
                    "PEPE": ["pepe"],
                    "TON": ["ton", "toncoin"],
                }
                terms = aliases.get(news_coin_filter.upper(), [kw])
                all_articles = [
                    a
                    for a in all_articles
                    if any(
                        tm in a["title"].lower() or tm in a["body"].lower()
                        for tm in terms
                    )
                ]

            if not all_articles:
                st.info(
                    f"No news found"
                    f"{' for ' + news_coin_filter if news_coin_filter != 'All' else ''}"
                    f"{' in ' + news_time_filter.lower() if news_time_filter != 'Latest' else ''}. "
                    "Try 'All' coins or a wider time range."
                )
                return

            st.markdown(
                f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:14px;">'
                f'<span style="display:inline-block;width:8px;height:8px;border-radius:50%;'
                f'background:{t["green"]};animation:alphalens-live-pulse 2s ease-in-out infinite;"></span>'
                f'<span style="font-size:0.78rem;color:{t["label"]};font-weight:500;">'
                f"{len(all_articles)} articles"
                f" · {len(st.session_state['news_archive'])} in archive</span>"
                f"</div>",
                unsafe_allow_html=True,
            )

            for article in all_articles:
                _title = html.escape(article["title"])
                _source = html.escape(article["source_name"])
                _url = article["url"]
                _body = html.escape(article["body"])
                _pub = article["published_on"]

                # Time ago
                if _pub <= 0:
                    _tago = ""
                else:
                    diff = now - _pub
                    if diff < 60:
                        _tago = "just now"
                    elif diff < 3600:
                        _tago = f"{int(diff / 60)}m ago"
                    elif diff < 86400:
                        _tago = f"{int(diff / 3600)}h ago"
                    else:
                        d = int(diff / 86400)
                        h = int((diff % 86400) / 3600)
                        _tago = f"{d}d {h}h ago"

                st.markdown(
                    f'<div class="alphalens-news-card">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'align-items:center;margin-bottom:8px;">'
                    f'<div style="display:flex;align-items:center;gap:10px;">'
                    f'<span class="news-source">{_source}</span>'
                    f'<span class="news-time">{_tago}</span>'
                    f"</div>"
                    f'<span class="news-watermark">RSS</span>'
                    f"</div>"
                    f'<a href="{_url}" target="_blank" style="text-decoration:none;">'
                    f'<div class="news-title">{_title}</div></a>'
                    f'<div class="news-body">{_body}{"..." if len(_body) >= 297 else ""}</div>'
                    f'<div class="news-footer">'
                    f'<span class="news-watermark">{_source}</span>'
                    f"</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

        _news_feed()
